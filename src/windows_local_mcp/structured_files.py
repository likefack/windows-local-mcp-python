"""Closed-world, broker-direct structured file transformations.

This module intentionally accepts declarative data only. It never evaluates user-supplied code,
loads a project-controlled helper, or invokes Office. The caller owns workspace path validation,
transactional replacement, and audit. A future external helper may be broker-directed when WLMCP
can close and verify its effective side effects; otherwise it belongs in Codex Sandbox. In either
case it should return an artifact to this boundary rather than writing the workspace itself.
"""

from __future__ import annotations

import base64
import csv
import io
import math
import posixpath
import re
import warnings
import zipfile
from copy import copy
from dataclasses import dataclass
from pathlib import PurePosixPath, PureWindowsPath
from typing import Any
from xml.etree import ElementTree

from .config import Settings
from .util import sha256_bytes

_FORMATS = {"docx", "xlsx", "csv", "tsv", "zip", "image"}
_EXTENSIONS = {
    "docx": {".docx"},
    "xlsx": {".xlsx"},
    "csv": {".csv"},
    "tsv": {".tsv", ".tab"},
    "zip": {".zip"},
    "image": {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"},
}
_ZIP_RESERVED = {"CON", "PRN", "AUX", "NUL", *(f"COM{i}" for i in range(1, 10)), *(f"LPT{i}" for i in range(1, 10))}

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"
_SHEET_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


class StructuredFileError(ValueError):
    """A safe, user-actionable rejection of a structured file request."""


def infer_format(path: str, requested: str | None = None) -> str:
    suffix = PureWindowsPath(path).suffix.casefold()
    if requested is not None:
        kind = requested.casefold()
        if kind not in _FORMATS:
            raise StructuredFileError("format must be docx, xlsx, csv, tsv, zip, or image")
        if suffix not in _EXTENSIONS[kind]:
            raise StructuredFileError(f"path extension is not supported for {kind}")
        return kind
    for kind, extensions in _EXTENSIONS.items():
        if suffix in extensions:
            return kind
    raise StructuredFileError("unsupported structured file extension")


def _require_size(data: bytes, settings: Settings) -> None:
    if len(data) > settings.max_structured_file_bytes:
        raise StructuredFileError("structured file exceeds max_structured_file_bytes")


def _dependency(name: str) -> None:
    raise StructuredFileError(
        f"{name} support is unavailable; install the Windows Local MCP structured-file dependencies"
    )


def _docx_modules() -> tuple[Any, Any, Any, Any]:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        _dependency("DOCX")
    return Document, WD_ORIENT, WD_ALIGN_PARAGRAPH, (Inches, Pt, RGBColor)


def _xlsx_modules() -> tuple[Any, Any, Any, Any, Any, Any, Any, Any, Any]:
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.chart import BarChart, LineChart, Reference
        from openpyxl.formatting.rule import CellIsRule
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.worksheet.datavalidation import DataValidation
    except ImportError:
        _dependency("XLSX")
    return (
        Workbook,
        load_workbook,
        BarChart,
        LineChart,
        Reference,
        CellIsRule,
        (Alignment, Border, Font, PatternFill, Side),
        DataValidation,
        None,
    )


def _image_module() -> Any:
    try:
        from PIL import Image, UnidentifiedImageError
    except ImportError:
        _dependency("image")
    return Image, UnidentifiedImageError


def _operation(raw: Any) -> dict[str, Any]:
    if not isinstance(raw, dict) or not isinstance(raw.get("op"), str):
        raise StructuredFileError("each operation must be an object with a string op")
    return raw


def _text(value: Any, name: str = "text") -> str:
    if not isinstance(value, str):
        raise StructuredFileError(f"{name} must be a string")
    return value


def _index(value: Any, name: str, *, minimum: int = 0) -> int:
    if isinstance(value, bool) or not isinstance(value, int) or value < minimum:
        raise StructuredFileError(f"{name} must be an integer >= {minimum}")
    return value


def _bounded(items: list[Any], settings: Settings, label: str) -> None:
    if len(items) > settings.max_structured_elements:
        raise StructuredFileError(f"{label} exceeds max_structured_elements")


def _require_package_bounds(
    archive: zipfile.ZipFile, settings: Settings, label: str
) -> list[zipfile.ZipInfo]:
    infos = archive.infolist()
    if len(infos) > settings.max_zip_entries:
        raise StructuredFileError(f"{label} package exceeds max_zip_entries")
    expanded = 0
    for info in infos:
        if info.flag_bits & 0x1:
            raise StructuredFileError(f"encrypted {label} package parts are unsupported")
        expanded += info.file_size
        if expanded > settings.max_zip_expanded_bytes:
            raise StructuredFileError(f"{label} package exceeds max_zip_expanded_bytes")
    return infos


def _parse_package_xml(data: bytes, label: str) -> ElementTree.Element:
    try:
        return ElementTree.fromstring(data)
    except ElementTree.ParseError as error:
        raise StructuredFileError(f"invalid {label} XML") from error


def _serialize_package_xml(root: ElementTree.Element) -> bytes:
    return ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)


def _rewrite_package(data: bytes, replacements: dict[str, bytes]) -> bytes:
    """Rewrite selected OPC parts while preserving every other ZIP member payload and metadata."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as source:
            infos = source.infolist()
            names = [info.filename for info in infos]
            for name in replacements:
                if names.count(name) != 1:
                    raise StructuredFileError(f"structured package part is missing or duplicated: {name}")
            output = io.BytesIO()
            with zipfile.ZipFile(output, "w") as destination:
                destination.comment = source.comment
                for info in infos:
                    destination.writestr(info, replacements.get(info.filename, source.read(info)))
            return output.getvalue()
    except zipfile.BadZipFile as error:
        raise StructuredFileError("invalid structured ZIP package") from error


def _replace_xml_text_nodes(
    paragraph: ElementTree.Element, search: str, replacement: str
) -> int:
    blocked = {"ins", "del", "moveFrom", "moveTo", "fldSimple", "sdt"}
    nodes: list[ElementTree.Element] = []

    def collect(element: ElementTree.Element, blocked_ancestor: bool = False) -> None:
        local_name = element.tag.rsplit("}", 1)[-1]
        blocked_here = blocked_ancestor or local_name in blocked
        if element.tag == f"{{{_WORD_NS}}}t" and not blocked_here:
            nodes.append(element)
            return
        for child in element:
            collect(child, blocked_here)

    collect(paragraph)
    joined = "".join(node.text or "" for node in nodes)
    positions: list[int] = []
    start = 0
    while True:
        found = joined.find(search, start)
        if found < 0:
            break
        positions.append(found)
        start = found + len(search)
    for found in reversed(positions):
        end = found + len(search)
        cursor = 0
        first_index = last_index = -1
        first_offset = last_offset = 0
        for index, node in enumerate(nodes):
            text = node.text or ""
            next_cursor = cursor + len(text)
            if first_index < 0 and found < next_cursor:
                first_index, first_offset = index, found - cursor
            if end <= next_cursor:
                last_index, last_offset = index, end - cursor
                break
            cursor = next_cursor
        if first_index < 0 or last_index < 0:
            continue
        if first_index == last_index:
            text = nodes[first_index].text or ""
            nodes[first_index].text = text[:first_offset] + replacement + text[last_offset:]
        else:
            first_text = nodes[first_index].text or ""
            last_text = nodes[last_index].text or ""
            nodes[first_index].text = first_text[:first_offset] + replacement
            for node in nodes[first_index + 1 : last_index]:
                node.text = ""
            nodes[last_index].text = last_text[last_offset:]
        for node in nodes[first_index : last_index + 1]:
            if (node.text or "").startswith(" ") or (node.text or "").endswith(" "):
                node.set(f"{{{_XML_NS}}}space", "preserve")
    return len(positions)


def _transform_docx_package_preserving(
    data: bytes, operations: list[Any], unsupported: list[str]
) -> tuple[bytes, dict[str, Any]]:
    if "digital signatures" in unsupported:
        raise StructuredFileError("DOCX write would invalidate digital signatures")
    parsed_operations = [_operation(raw) for raw in operations]
    allowed = {"replace_text", "metadata_set"}
    rejected = sorted({op["op"] for op in parsed_operations} - allowed)
    if rejected:
        raise StructuredFileError(
            "DOCX contains features that require package-preserving edits; supported operations are "
            "replace_text and metadata_set (rejected: " + ", ".join(rejected) + ")"
        )
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names = archive.namelist()
            text_parts = [
                name
                for name in names
                if name.casefold() == "word/document.xml"
                or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name, flags=re.IGNORECASE)
            ]
            replacements: dict[str, bytes] = {}
            text_roots = {
                name: _parse_package_xml(archive.read(name), f"DOCX part {name}")
                for name in text_parts
            }
            core_name = next(
                (name for name in names if name.casefold() == "docprops/core.xml"), None
            )
            core_root = (
                _parse_package_xml(archive.read(core_name), "DOCX core properties")
                if core_name is not None
                else None
            )
    except zipfile.BadZipFile as error:
        raise StructuredFileError("invalid DOCX package") from error

    core_namespaces = {
        "title": "{http://purl.org/dc/elements/1.1/}title",
        "subject": "{http://purl.org/dc/elements/1.1/}subject",
        "author": "{http://purl.org/dc/elements/1.1/}creator",
        "keywords": "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}keywords",
        "comments": "{http://purl.org/dc/elements/1.1/}description",
        "category": "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}category",
    }
    text_changed = False
    core_changed = False
    for op in parsed_operations:
        if op["op"] == "replace_text":
            search = _text(op.get("search"), "search")
            replacement = _text(op.get("replace"), "replace")
            if not search:
                raise StructuredFileError("search must not be empty")
            count = 0
            for root in text_roots.values():
                for paragraph in root.iter(f"{{{_WORD_NS}}}p"):
                    count += _replace_xml_text_nodes(paragraph, search, replacement)
            text_changed = text_changed or count > 0
            if op.get("require_match", False) and count == 0:
                raise StructuredFileError("replace_text found no safe package-preserving match")
        else:
            if core_name is None or core_root is None:
                raise StructuredFileError("DOCX core properties part is unavailable")
            values = op.get("values")
            if not isinstance(values, dict):
                raise StructuredFileError("metadata values must be an object")
            for key, value in values.items():
                tag = core_namespaces.get(key)
                if tag is None:
                    raise StructuredFileError(f"unsupported metadata field: {key}")
                element = core_root.find(tag)
                if element is None:
                    element = ElementTree.SubElement(core_root, tag)
                element.text = _text(value, key)
                core_changed = True
    if text_changed:
        replacements.update(
            {name: _serialize_package_xml(root) for name, root in text_roots.items()}
        )
    if core_changed and core_name is not None and core_root is not None:
        replacements[core_name] = _serialize_package_xml(core_root)
    output = _rewrite_package(data, replacements) if replacements else data
    return output, {
        "preservation_mode": "package_patch",
        "preserved_unsupported_features": unsupported,
    }


def _docx_has_unsupported_features(data: bytes, settings: Settings) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = _require_package_bounds(archive, settings, "DOCX")
            archive_names = [info.filename for info in infos]
            names = {name.casefold() for name in archive_names}
            document = archive.read("word/document.xml")
            word_xml = [
                archive.read(name)
                for name in archive_names
                if name.casefold().startswith("word/")
                and name.casefold().endswith(".xml")
            ]
    except (KeyError, zipfile.BadZipFile) as error:
        raise StructuredFileError("invalid DOCX package") from error
    found: list[str] = []
    if any("vbaproject.bin" in name for name in names):
        found.append("VBA/macros")
    if any("activex" in name for name in names):
        found.append("ActiveX")
    if any("diagrams/" in name for name in names):
        found.append("SmartArt")
    if any("embeddings/" in name for name in names):
        found.append("embedded objects")
    if any("afchunk" in name for name in names):
        found.append("alternative-format chunks")
    for marker, label in (
        ("word/comments", "comments"),
        ("word/footnotes.xml", "footnotes"),
        ("word/endnotes.xml", "endnotes"),
        ("word/glossary/", "glossary document"),
    ):
        if any(marker in name for name in names):
            found.append(label)
    if any("_xmlsignatures/" in name for name in names):
        found.append("digital signatures")
    if b"<w:ins" in document or b"<w:del" in document or b"<w:move" in document:
        found.append("tracked changes")
    if any(b":dataBinding" in xml for xml in word_xml):
        found.append("custom XML data binding")
    return found


def _paragraph_info(paragraph: Any, index: int) -> dict[str, Any]:
    return {
        "index": index,
        "text": paragraph.text,
        "style": getattr(paragraph.style, "name", None),
        "runs": [
            {"index": run_index, "text": run.text, "bold": run.bold, "italic": run.italic}
            for run_index, run in enumerate(paragraph.runs)
        ],
    }


def _length_inches(value: Any) -> float | None:
    return None if value is None else float(value.inches)


def _inspect_docx(data: bytes, settings: Settings) -> dict[str, Any]:
    Document, _, _, _ = _docx_modules()
    unsupported = _docx_has_unsupported_features(data, settings)
    document = Document(io.BytesIO(data))
    _bounded(document.paragraphs, settings, "DOCX paragraphs")
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        cells = sum(len(row.cells) for row in table.rows)
        if cells > settings.max_structured_elements:
            raise StructuredFileError("DOCX table exceeds max_structured_elements")
        tables.append(
            {
                "index": table_index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": [[cell.text for cell in row.cells] for row in table.rows[:20]],
            }
        )
    sections = []
    for section_index, section in enumerate(document.sections):
        sections.append(
            {
                "index": section_index,
                "orientation": str(section.orientation),
                "page_width_inches": _length_inches(section.page_width),
                "page_height_inches": _length_inches(section.page_height),
                "margins_inches": {
                    "top": _length_inches(section.top_margin),
                    "bottom": _length_inches(section.bottom_margin),
                    "left": _length_inches(section.left_margin),
                    "right": _length_inches(section.right_margin),
                },
                "header": [paragraph.text for paragraph in section.header.paragraphs[:50]],
                "footer": [paragraph.text for paragraph in section.footer.paragraphs[:50]],
            }
        )
    props = document.core_properties
    return {
        "format": "docx",
        "paragraph_count": len(document.paragraphs),
        "paragraphs": [_paragraph_info(p, i) for i, p in enumerate(document.paragraphs[:200])],
        "table_count": len(document.tables),
        "tables": tables[:50],
        "section_count": len(document.sections),
        "sections": sections,
        "metadata": {key: getattr(props, key) for key in ("title", "subject", "author", "keywords", "comments")},
        "write_rejected_features": unsupported,
        "package_patch_supported_operations": (
            [] if "digital signatures" in unsupported else ["replace_text", "metadata_set"]
        ),
        "truncated": len(document.paragraphs) > 200 or len(document.tables) > 50,
    }


def _apply_docx_run_format(run: Any, spec: dict[str, Any]) -> None:
    if not isinstance(spec, dict):
        raise StructuredFileError("run formatting must be an object")
    _, _, _, units = _docx_modules()
    _, Pt, RGBColor = units
    if "font_name" in spec:
        run.font.name = _text(spec["font_name"], "font_name")
    if "font_size_pt" in spec:
        size = spec["font_size_pt"]
        if not isinstance(size, (int, float)) or size <= 0:
            raise StructuredFileError("font_size_pt must be a positive number")
        run.font.size = Pt(size)
    for key, attr in (("bold", "bold"), ("italic", "italic"), ("underline", "underline")):
        if key in spec:
            if not isinstance(spec[key], bool):
                raise StructuredFileError(f"{key} must be boolean")
            setattr(run.font, attr, spec[key])
    if "color" in spec:
        color = _text(spec["color"], "color")
        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            raise StructuredFileError("color must be six hexadecimal digits")
        run.font.color.rgb = RGBColor.from_string(color.upper())


def _apply_docx_format(paragraph: Any, spec: dict[str, Any]) -> None:
    _, _, WD_ALIGN_PARAGRAPH, units = _docx_modules()
    Inches, Pt, _ = units
    if "style" in spec:
        paragraph.style = _text(spec["style"], "style")
    fmt = paragraph.paragraph_format
    if "alignment" in spec:
        value = _text(spec["alignment"], "alignment").upper()
        if not hasattr(WD_ALIGN_PARAGRAPH, value):
            raise StructuredFileError("unsupported paragraph alignment")
        paragraph.alignment = getattr(WD_ALIGN_PARAGRAPH, value)
    for key, attr in (("left_indent_inches", "left_indent"), ("right_indent_inches", "right_indent"), ("first_line_indent_inches", "first_line_indent"), ("space_before_pt", "space_before"), ("space_after_pt", "space_after")):
        if key in spec:
            value = spec[key]
            if not isinstance(value, (int, float)):
                raise StructuredFileError(f"{key} must be a number")
            setattr(fmt, attr, Inches(value) if "indent" in key else Pt(value))
    if "line_spacing" in spec:
        value = spec["line_spacing"]
        if not isinstance(value, (int, float)):
            raise StructuredFileError("line_spacing must be a number")
        fmt.line_spacing = value
    run_spec = spec.get("run")
    if run_spec is not None:
        for run in paragraph.runs:
            _apply_docx_run_format(run, run_spec)


def _docx_paragraphs(document: Any) -> list[Any]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def _paragraph_has_complex_inline(paragraph: Any) -> bool:
    xml = paragraph._p.xml
    return any(marker in xml for marker in ("<w:hyperlink", "<w:drawing", "<w:object", "<w:fldChar", "<w:instrText"))


def _cell_has_complex_inline(cell: Any) -> bool:
    return any(_paragraph_has_complex_inline(paragraph) for paragraph in cell.paragraphs)


def _require_docx_bounds(document: Any, settings: Settings) -> None:
    paragraphs = _docx_paragraphs(document)
    if len(paragraphs) > settings.max_structured_elements:
        raise StructuredFileError("DOCX paragraphs exceed max_structured_elements")
    cells = sum(len(row.cells) for table in document.tables for row in table.rows)
    if cells > settings.max_structured_elements:
        raise StructuredFileError("DOCX table cells exceed max_structured_elements")


def _replace_text_preserving_runs(paragraph: Any, search: str, replacement: str) -> int:
    runs = list(paragraph.runs)
    joined = "".join(run.text for run in runs)
    positions: list[int] = []
    start = 0
    while True:
        found = joined.find(search, start)
        if found < 0:
            break
        positions.append(found)
        start = found + len(search)
    for found in reversed(positions):
        end = found + len(search)
        cursor = 0
        first_index = last_index = -1
        first_offset = last_offset = 0
        for index, run in enumerate(runs):
            next_cursor = cursor + len(run.text)
            if first_index < 0 and found < next_cursor:
                first_index, first_offset = index, found - cursor
            if end <= next_cursor:
                last_index, last_offset = index, end - cursor
                break
            cursor = next_cursor
        if first_index < 0 or last_index < 0:
            continue
        if first_index == last_index:
            text = runs[first_index].text
            runs[first_index].text = text[:first_offset] + replacement + text[last_offset:]
        else:
            first_text = runs[first_index].text
            last_text = runs[last_index].text
            runs[first_index].text = first_text[:first_offset] + replacement
            for run in runs[first_index + 1 : last_index]:
                run.text = ""
            runs[last_index].text = last_text[last_offset:]
    return len(positions)


def _apply_docx_cell_format(cell: Any, spec: dict[str, Any]) -> None:
    if not isinstance(spec, dict):
        raise StructuredFileError("cell format must be an object")
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if "vertical_alignment" in spec:
        name = _text(spec["vertical_alignment"], "vertical_alignment").upper()
        if not hasattr(WD_CELL_VERTICAL_ALIGNMENT, name):
            raise StructuredFileError("unsupported cell vertical alignment")
        cell.vertical_alignment = getattr(WD_CELL_VERTICAL_ALIGNMENT, name)
    if "shading" in spec:
        color = _text(spec["shading"], "shading")
        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            raise StructuredFileError("cell shading must be six hexadecimal digits")
        properties = cell._tc.get_or_add_tcPr()
        for existing in properties.findall(qn("w:shd")):
            properties.remove(existing)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color.upper())
        properties.append(shading)
    if "paragraph" in spec:
        paragraph_spec = spec["paragraph"]
        if not isinstance(paragraph_spec, dict):
            raise StructuredFileError("cell paragraph format must be an object")
        for paragraph in cell.paragraphs:
            _apply_docx_format(paragraph, paragraph_spec)


def _transform_docx(
    data: bytes, operations: list[Any], settings: Settings
) -> tuple[bytes, dict[str, Any]]:
    unsupported = _docx_has_unsupported_features(data, settings)
    if unsupported:
        return _transform_docx_package_preserving(data, operations, unsupported)
    Document, WD_ORIENT, _, units = _docx_modules()
    Inches, _, _ = units
    document = Document(io.BytesIO(data))
    _require_docx_bounds(document, settings)
    for raw in operations:
        op = _operation(raw)
        name = op["op"]
        if name == "paragraph_append":
            paragraph = document.add_paragraph(_text(op.get("text", "")))
            _apply_docx_format(paragraph, op.get("format", {}))
        elif name == "paragraph_update":
            paragraph = document.paragraphs[_index(op.get("index"), "index")]
            if "text" in op:
                if _paragraph_has_complex_inline(paragraph) and not op.get("allow_inline_loss", False):
                    raise StructuredFileError(
                        "paragraph contains a hyperlink, field, image, or embedded object; "
                        "full-text replacement is unsupported without allow_inline_loss"
                    )
                # This explicit replacement intentionally replaces only the selected paragraph's runs.
                paragraph.clear()
                paragraph.add_run(_text(op["text"]))
            if "format" in op:
                if not isinstance(op["format"], dict):
                    raise StructuredFileError("format must be an object")
                _apply_docx_format(paragraph, op["format"])
        elif name == "paragraph_delete":
            paragraph = document.paragraphs[_index(op.get("index"), "index")]
            paragraph._element.getparent().remove(paragraph._element)
        elif name == "run_update":
            paragraph = document.paragraphs[_index(op.get("paragraph"), "paragraph")]
            run = paragraph.runs[_index(op.get("run"), "run")]
            if "text" in op:
                run.text = _text(op["text"])
            if "format" in op:
                _apply_docx_run_format(run, op["format"])
        elif name == "replace_text":
            search = _text(op.get("search"), "search")
            replacement = _text(op.get("replace"), "replace")
            if not search:
                raise StructuredFileError("search must not be empty")
            count = 0
            for paragraph in _docx_paragraphs(document):
                count += _replace_text_preserving_runs(paragraph, search, replacement)
            if op.get("require_match", False) and count == 0:
                raise StructuredFileError("replace_text found no match")
        elif name == "table_cell_set":
            table = document.tables[_index(op.get("table"), "table")]
            cell = table.cell(_index(op.get("row"), "row"), _index(op.get("column"), "column"))
            if _cell_has_complex_inline(cell) and not op.get("allow_inline_loss", False):
                raise StructuredFileError(
                    "table cell contains a hyperlink, field, image, or embedded object; "
                    "full-text replacement is unsupported without allow_inline_loss"
                )
            cell.text = _text(op.get("text", ""))
            if "format" in op:
                _apply_docx_cell_format(cell, op["format"])
        elif name == "table_row_add":
            table = document.tables[_index(op.get("table"), "table")]
            values = op.get("values", [])
            if not isinstance(values, list) or len(values) > len(table.columns):
                raise StructuredFileError("values must fit in the table columns")
            row = table.add_row()
            for cell, value in zip(row.cells, values, strict=False):
                cell.text = _text(value, "table value")
        elif name == "table_row_insert":
            table = document.tables[_index(op.get("table"), "table")]
            index = _index(op.get("row"), "row")
            if index > len(table.rows):
                raise StructuredFileError("table row index is outside the table")
            values = op.get("values", [])
            if not isinstance(values, list) or len(values) > len(table.columns):
                raise StructuredFileError("values must fit in the table columns")
            row = table.add_row()
            if index < len(table.rows) - 1:
                table.rows[index]._tr.addprevious(row._tr)
            for cell, value in zip(row.cells, values, strict=False):
                cell.text = _text(value, "table value")
        elif name == "table_row_delete":
            table = document.tables[_index(op.get("table"), "table")]
            if len(table.rows) <= 1:
                raise StructuredFileError("cannot delete the last table row")
            row = table.rows[_index(op.get("row"), "row")]
            row._tr.getparent().remove(row._tr)
        elif name == "table_column_add":
            table = document.tables[_index(op.get("table"), "table")]
            width = op.get("width_inches", 1.0)
            if not isinstance(width, (int, float)) or width <= 0:
                raise StructuredFileError("width_inches must be positive")
            column = table.add_column(Inches(width))
            values = op.get("values", [])
            if not isinstance(values, list) or len(values) > len(column.cells):
                raise StructuredFileError("column values must fit in the table rows")
            for index, cell in enumerate(column.cells):
                if index < len(values):
                    cell.text = _text(values[index], "table value")
        elif name == "table_column_delete":
            table = document.tables[_index(op.get("table"), "table")]
            if len(table.columns) <= 1:
                raise StructuredFileError("cannot delete the last table column")
            column = _index(op.get("column"), "column")
            if column >= len(table.columns):
                raise StructuredFileError("table column index is outside the table")
            for row in table.rows:
                row._tr.remove(row.cells[column]._tc)
            grid = table._tbl.tblGrid
            grid.remove(grid.gridCol_lst[column])
        elif name == "table_cell_format":
            table = document.tables[_index(op.get("table"), "table")]
            cell = table.cell(_index(op.get("row"), "row"), _index(op.get("column"), "column"))
            _apply_docx_cell_format(cell, op.get("format", {}))
        elif name == "table_add":
            rows = _index(op.get("rows"), "rows", minimum=1)
            columns = _index(op.get("columns"), "columns", minimum=1)
            if rows * columns > settings.max_structured_elements:
                raise StructuredFileError("new table exceeds max_structured_elements")
            table = document.add_table(rows=rows, cols=columns)
            if "style" in op:
                table.style = _text(op["style"], "style")
            values = op.get("values", [])
            if not isinstance(values, list):
                raise StructuredFileError("table values must be an array")
            for row_index, values_row in enumerate(values):
                if row_index >= rows or not isinstance(values_row, list) or len(values_row) > columns:
                    raise StructuredFileError("table values must fit the requested shape")
                for column_index, value in enumerate(values_row):
                    table.cell(row_index, column_index).text = _text(value, "table value")
        elif name in {"header_footer_append", "header_footer_set"}:
            section = document.sections[_index(op.get("section", 0), "section")]
            area = _text(op.get("area", "header"), "area")
            if area not in {"header", "footer"}:
                raise StructuredFileError("area must be header or footer")
            container = getattr(section, area)
            if name == "header_footer_append":
                container.add_paragraph(_text(op.get("text", "")))
            else:
                index = _index(op.get("paragraph", 0), "paragraph")
                if index >= len(container.paragraphs):
                    raise StructuredFileError("header/footer paragraph index is outside the section")
                paragraph = container.paragraphs[index]
                if _paragraph_has_complex_inline(paragraph) and not op.get("allow_inline_loss", False):
                    raise StructuredFileError("header/footer paragraph contains unsupported inline content")
                paragraph.clear()
                paragraph.add_run(_text(op.get("text", "")))
                if "format" in op:
                    _apply_docx_format(paragraph, op["format"])
        elif name == "style_update":
            style = document.styles[_text(op.get("style"), "style")]
            if "run" in op:
                _apply_docx_run_format(style, op["run"])
            paragraph_values = op.get("paragraph")
            if paragraph_values is not None:
                if not isinstance(paragraph_values, dict):
                    raise StructuredFileError("style paragraph format must be an object")
                _, _, alignments, _style_units = _docx_modules()
                style_format = style.paragraph_format
                if "alignment" in paragraph_values:
                    value = _text(paragraph_values["alignment"], "alignment").upper()
                    if not hasattr(alignments, value):
                        raise StructuredFileError("unsupported paragraph alignment")
                    style_format.alignment = getattr(alignments, value)
                for key, attr in (
                    ("left_indent_inches", "left_indent"),
                    ("right_indent_inches", "right_indent"),
                    ("first_line_indent_inches", "first_line_indent"),
                    ("space_before_pt", "space_before"),
                    ("space_after_pt", "space_after"),
                ):
                    if key in paragraph_values:
                        value = paragraph_values[key]
                        if not isinstance(value, (int, float)):
                            raise StructuredFileError(f"{key} must be a number")
                        setattr(
                            style_format,
                            attr,
                            Inches(value) if "indent" in key else units[1](value),
                        )
        elif name == "metadata_set":
            values = op.get("values")
            if not isinstance(values, dict):
                raise StructuredFileError("metadata values must be an object")
            allowed = {"title", "subject", "author", "keywords", "comments", "category"}
            for key, value in values.items():
                if key not in allowed:
                    raise StructuredFileError(f"unsupported metadata field: {key}")
                setattr(document.core_properties, key, _text(value, key))
        elif name == "section_set":
            section = document.sections[_index(op.get("section", 0), "section")]
            if "orientation" in op:
                orientation = _text(op["orientation"], "orientation").upper()
                if orientation not in {"PORTRAIT", "LANDSCAPE"}:
                    raise StructuredFileError("orientation must be portrait or landscape")
                section.orientation = getattr(WD_ORIENT, orientation)
                section.page_width, section.page_height = section.page_height, section.page_width
            for key, attr in (("page_width_inches", "page_width"), ("page_height_inches", "page_height"), ("top_margin_inches", "top_margin"), ("bottom_margin_inches", "bottom_margin"), ("left_margin_inches", "left_margin"), ("right_margin_inches", "right_margin")):
                if key in op:
                    value = op[key]
                    if not isinstance(value, (int, float)) or value < 0:
                        raise StructuredFileError(f"{key} must be a non-negative number")
                    setattr(section, attr, Inches(value))
        else:
            raise StructuredFileError(f"unsupported DOCX operation: {name}")
        _require_docx_bounds(document, settings)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), {"preservation_mode": "library_rewrite"}


def _xlsx_unsupported(data: bytes, settings: Settings) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = _require_package_bounds(archive, settings, "XLSX")
            archive_names = [info.filename for info in infos]
            names = {name.casefold() for name in archive_names}
            xml_parts = [
                archive.read(name)
                for name in archive_names
                if name.casefold().endswith(".xml")
            ]
    except zipfile.BadZipFile as error:
        raise StructuredFileError("invalid XLSX package") from error
    blocked = []
    for marker, label in (
        ("vbaproject.bin", "VBA/macros"),
        ("activex", "ActiveX"),
        ("connections.xml", "external data connections"),
        ("slicer", "slicers"),
        ("threadedcomments", "threaded comments"),
        ("persons/", "threaded-comment persons"),
        ("richdata/", "rich data"),
        ("ctrlprops/", "form controls"),
        ("pivottable", "pivot tables"),
        ("pivotcache", "pivot caches"),
        ("embeddings/", "embedded objects"),
        ("customxml/", "custom XML"),
        ("_xmlsignatures/", "digital signatures"),
    ):
        if any(marker in name for name in names):
            blocked.append(label)
    if any(re.search(rb"<(?:[A-Za-z0-9_]+:)?extLst\b", part) for part in xml_parts):
        blocked.append("unsupported extension lists")
    return blocked


def _xlsx_sheet_part(
    workbook: ElementTree.Element,
    relationships: ElementTree.Element,
    sheet_name: str,
    package_names: dict[str, str],
) -> str:
    relation_id: str | None = None
    for sheet in workbook.iter(f"{{{_SHEET_NS}}}sheet"):
        if sheet.get("name") == sheet_name:
            relation_id = sheet.get(f"{{{_OFFICE_REL_NS}}}id")
            break
    if relation_id is None:
        raise StructuredFileError(f"XLSX worksheet does not exist: {sheet_name}")
    target: str | None = None
    for relation in relationships.iter(f"{{{_PACKAGE_REL_NS}}}Relationship"):
        if relation.get("Id") == relation_id:
            if str(relation.get("TargetMode", "")).casefold() == "external":
                raise StructuredFileError("XLSX worksheet relationship is external")
            target = relation.get("Target")
            break
    if not target:
        raise StructuredFileError("XLSX worksheet relationship is missing")
    normalized = (
        posixpath.normpath(target.lstrip("/"))
        if target.startswith("/")
        else posixpath.normpath(posixpath.join("xl", target))
    )
    if PurePosixPath(normalized).is_absolute() or not normalized.casefold().startswith("xl/"):
        raise StructuredFileError("XLSX worksheet relationship escapes the package")
    actual = package_names.get(normalized.casefold())
    if actual is None:
        raise StructuredFileError("XLSX worksheet part is missing")
    return actual


def _xlsx_cell_payload(cell: ElementTree.Element, value: Any) -> None:
    formula = cell.find(f"{{{_SHEET_NS}}}f")
    if formula is not None and formula.get("t") in {"shared", "array", "dataTable"}:
        raise StructuredFileError("cell belongs to a grouped formula and cannot be patched locally")
    removable = {
        f"{{{_SHEET_NS}}}f",
        f"{{{_SHEET_NS}}}v",
        f"{{{_SHEET_NS}}}is",
    }
    for child in list(cell):
        if child.tag in removable:
            cell.remove(child)

    cell.attrib.pop("t", None)
    payload: ElementTree.Element | None = None
    if value is None:
        return
    if isinstance(value, bool):
        cell.set("t", "b")
        payload = ElementTree.Element(f"{{{_SHEET_NS}}}v")
        payload.text = "1" if value else "0"
    elif isinstance(value, (int, float)) and not isinstance(value, bool):
        if isinstance(value, float) and not math.isfinite(value):
            raise StructuredFileError("XLSX numeric values must be finite")
        payload = ElementTree.Element(f"{{{_SHEET_NS}}}v")
        payload.text = str(value)
    elif isinstance(value, str) and value.startswith("="):
        if len(value) > 8193:
            raise StructuredFileError("XLSX formula exceeds the supported length")
        payload = ElementTree.Element(f"{{{_SHEET_NS}}}f")
        payload.text = value[1:]
    elif isinstance(value, str):
        if len(value) > 32767:
            raise StructuredFileError("XLSX text exceeds the cell length limit")
        cell.set("t", "inlineStr")
        inline = ElementTree.Element(f"{{{_SHEET_NS}}}is")
        payload = ElementTree.SubElement(inline, f"{{{_SHEET_NS}}}t")
        payload.text = value
        if value.startswith(" ") or value.endswith(" "):
            payload.set(f"{{{_XML_NS}}}space", "preserve")
        payload = inline
    else:
        raise StructuredFileError("XLSX cell values must be string, number, boolean, or null")

    extension = cell.find(f"{{{_SHEET_NS}}}extLst")
    if extension is None:
        cell.append(payload)
    else:
        cell.insert(list(cell).index(extension), payload)


def _xlsx_dimension_bounds(root: ElementTree.Element) -> tuple[int, int, int, int]:
    dimension = root.find(f"{{{_SHEET_NS}}}dimension")
    if dimension is None or not dimension.get("ref"):
        return 1, 1, 1, 1
    try:
        from openpyxl.utils.cell import range_boundaries

        bounds = range_boundaries(str(dimension.get("ref")))
    except (TypeError, ValueError) as error:
        raise StructuredFileError("invalid XLSX worksheet dimension") from error
    if None in bounds:
        raise StructuredFileError("invalid XLSX worksheet dimension")
    return tuple(int(item) for item in bounds)  # type: ignore[return-value]


def _xlsx_patch_state(root: ElementTree.Element) -> dict[str, Any]:
    sheet_data = root.find(f"{{{_SHEET_NS}}}sheetData")
    if sheet_data is None:
        raise StructuredFileError("XLSX worksheet has no sheetData")
    from openpyxl.utils.cell import range_boundaries

    rows: dict[int, ElementTree.Element] = {}
    cells: dict[int, dict[int, ElementTree.Element]] = {}
    for row_element in sheet_data.findall(f"{{{_SHEET_NS}}}row"):
        try:
            row_number = int(str(row_element.get("r")))
        except (TypeError, ValueError) as error:
            raise StructuredFileError("XLSX worksheet row has an invalid index") from error
        if row_number in rows:
            raise StructuredFileError("XLSX worksheet contains duplicate rows")
        rows[row_number] = row_element
        row_cells: dict[int, ElementTree.Element] = {}
        for cell in row_element.findall(f"{{{_SHEET_NS}}}c"):
            reference = cell.get("r")
            if not reference:
                raise StructuredFileError("XLSX worksheet cell has no reference")
            try:
                column = range_boundaries(reference)[0]
            except (TypeError, ValueError) as error:
                raise StructuredFileError("XLSX worksheet cell has an invalid reference") from error
            if column in row_cells:
                raise StructuredFileError("XLSX worksheet contains duplicate cells")
            row_cells[column] = cell
        cells[row_number] = row_cells
    merged = []
    for item in root.iter(f"{{{_SHEET_NS}}}mergeCell"):
        reference = item.get("ref")
        if reference:
            merged.append(range_boundaries(reference))
    return {
        "root": root,
        "sheet_data": sheet_data,
        "rows": rows,
        "cells": cells,
        "merged": merged,
        "bounds": list(_xlsx_dimension_bounds(root)),
        "cell_count": sum(len(items) for items in cells.values()),
    }


def _xlsx_finalize_patch_state(state: dict[str, Any]) -> None:
    from openpyxl.utils.cell import get_column_letter

    sheet_data = state["sheet_data"]
    for row_number, row_element in state["rows"].items():
        cells = state["cells"][row_number]
        non_cells = [
            child for child in row_element if child.tag != f"{{{_SHEET_NS}}}c"
        ]
        row_element[:] = [cells[column] for column in sorted(cells)] + non_cells
    sheet_data[:] = sorted(sheet_data, key=lambda item: int(str(item.get("r", "0"))))
    min_column, min_row, max_column, max_row = state["bounds"]
    dimension = state["root"].find(f"{{{_SHEET_NS}}}dimension")
    if dimension is not None:
        dimension.set(
            "ref",
            f"{get_column_letter(min_column)}{min_row}:"
            f"{get_column_letter(max_column)}{max_row}",
        )


def _xlsx_patch_cell(
    state: dict[str, Any],
    bounds: tuple[int, int, int, int],
    value: Any,
    settings: Settings,
) -> None:
    column, row, max_column, max_row = bounds
    if column != max_column or row != max_row:
        raise StructuredFileError("cell must identify exactly one XLSX cell")
    current_min_column, current_min_row, current_max_column, current_max_row = state["bounds"]
    projected_max_column = max(current_max_column, column)
    projected_max_row = max(current_max_row, row)
    if projected_max_column * projected_max_row > settings.max_structured_elements:
        raise StructuredFileError("XLSX edit would exceed max_structured_elements")

    for merged_bounds in state["merged"]:
        min_col, min_row, max_col, merged_max_row = merged_bounds
        if (
            min_col <= column <= max_col
            and min_row <= row <= merged_max_row
            and (column, row) != (min_col, min_row)
        ):
            raise StructuredFileError("cannot patch a non-anchor cell inside a merged range")

    sheet_data = state["sheet_data"]
    row_element = state["rows"].get(row)
    if row_element is None:
        row_element = ElementTree.Element(f"{{{_SHEET_NS}}}row", {"r": str(row)})
        sheet_data.append(row_element)
        state["rows"][row] = row_element
        state["cells"][row] = {}
    from openpyxl.utils.cell import get_column_letter

    canonical_reference = f"{get_column_letter(column)}{row}"
    cell = state["cells"][row].get(column)
    if cell is None:
        cell = ElementTree.Element(f"{{{_SHEET_NS}}}c", {"r": canonical_reference})
        row_element.append(cell)
        state["cells"][row][column] = cell
        state["cell_count"] += 1
        if state["cell_count"] > settings.max_structured_elements:
            raise StructuredFileError("XLSX cells exceed max_structured_elements")
    _xlsx_cell_payload(cell, value)
    state["bounds"] = [
        min(current_min_column, column),
        min(current_min_row, row),
        projected_max_column,
        projected_max_row,
    ]


def _transform_xlsx_package_preserving(
    data: bytes,
    operations: list[Any],
    settings: Settings,
    unsupported: list[str],
) -> tuple[bytes, dict[str, Any]]:
    if "digital signatures" in unsupported:
        raise StructuredFileError("XLSX write would invalidate digital signatures")
    parsed_operations = [_operation(raw) for raw in operations]
    allowed = {"cell_set", "range_set", "range_clear"}
    rejected = sorted({op["op"] for op in parsed_operations} - allowed)
    if rejected:
        raise StructuredFileError(
            "XLSX contains features that require package-preserving edits; supported operations are "
            "cell_set, range_set, and range_clear (rejected: " + ", ".join(rejected) + ")"
        )
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names: dict[str, str] = {}
            for name in archive.namelist():
                folded = name.casefold()
                if folded in names:
                    raise StructuredFileError("XLSX package contains duplicate part names")
                names[folded] = name
            workbook_name = names.get("xl/workbook.xml")
            relationships_name = names.get("xl/_rels/workbook.xml.rels")
            if workbook_name is None or relationships_name is None:
                raise StructuredFileError("XLSX workbook relationships are missing")
            workbook = _parse_package_xml(archive.read(workbook_name), "XLSX workbook")
            relationships = _parse_package_xml(
                archive.read(relationships_name), "XLSX workbook relationships"
            )
            roots: dict[str, ElementTree.Element] = {}
            all_sheet_names = [
                str(sheet.get("name"))
                for sheet in workbook.iter(f"{{{_SHEET_NS}}}sheet")
                if sheet.get("name") is not None
            ]
            for sheet_name in all_sheet_names:
                part_name = _xlsx_sheet_part(workbook, relationships, sheet_name, names)
                if part_name not in roots:
                    roots[part_name] = _parse_package_xml(
                        archive.read(part_name), f"XLSX worksheet {sheet_name}"
                    )
    except zipfile.BadZipFile as error:
        raise StructuredFileError("invalid XLSX package") from error

    states = {name: _xlsx_patch_state(root) for name, root in roots.items()}
    if sum(int(state["cell_count"]) for state in states.values()) > settings.max_structured_elements:
        raise StructuredFileError("XLSX cells exceed max_structured_elements")
    changed_parts: set[str] = set()
    for op in parsed_operations:
        sheet_name = _text(op.get("sheet"), "sheet")
        part_name = _xlsx_sheet_part(workbook, relationships, sheet_name, names)
        state = states[part_name]
        changed_parts.add(part_name)
        if op["op"] == "cell_set":
            _reference, bounds = _xlsx_range_bounds(op.get("cell"), settings, "cell")
            _xlsx_patch_cell(state, bounds, op.get("value"), settings)
            continue
        _reference, bounds = _xlsx_range_bounds(op.get("range"), settings)
        min_column, min_row, max_column, max_row = bounds
        values = op.get("values")
        if op["op"] == "range_set":
            if not isinstance(values, list) or not all(isinstance(row, list) for row in values):
                raise StructuredFileError("values must be a two-dimensional array")
            if len(values) != max_row - min_row + 1 or any(
                len(value_row) != max_column - min_column + 1 for value_row in values
            ):
                raise StructuredFileError("values shape must match range")
        for row_offset, row in enumerate(range(min_row, max_row + 1)):
            for column_offset, column in enumerate(range(min_column, max_column + 1)):
                cell_value = (
                    values[row_offset][column_offset] if op["op"] == "range_set" else None
                )
                _xlsx_patch_cell(
                    state, (column, row, column, row), cell_value, settings
                )
    if sum(int(state["cell_count"]) for state in states.values()) > settings.max_structured_elements:
        raise StructuredFileError("XLSX cells exceed max_structured_elements")
    for part_name in changed_parts:
        _xlsx_finalize_patch_state(states[part_name])
    replacements = {
        name: _serialize_package_xml(roots[name]) for name in changed_parts
    }
    return _rewrite_package(data, replacements), {
        "preservation_mode": "package_patch",
        "preserved_unsupported_features": unsupported,
    }


def _sheet_cell_count(sheet: Any) -> int:
    return max(1, sheet.max_row) * max(1, sheet.max_column)


def _xlsx_range_bounds(
    value: Any, settings: Settings, label: str = "range"
) -> tuple[str, tuple[int, int, int, int]]:
    from openpyxl.utils.cell import range_boundaries

    reference = _text(value, label)
    try:
        min_col, min_row, max_col, max_row = range_boundaries(reference)
    except (TypeError, ValueError) as error:
        raise StructuredFileError(f"invalid XLSX {label}") from error
    if None in {min_col, min_row, max_col, max_row}:
        raise StructuredFileError(f"XLSX {label} must be an explicit A1 range")
    assert min_col is not None and min_row is not None
    assert max_col is not None and max_row is not None
    if min_col < 1 or min_row < 1 or max_col > 16384 or max_row > 1048576:
        raise StructuredFileError(f"XLSX {label} is outside worksheet limits")
    cells = (max_col - min_col + 1) * (max_row - min_row + 1)
    if cells > settings.max_structured_elements:
        raise StructuredFileError(f"XLSX {label} exceeds max_structured_elements")
    return reference, (min_col, min_row, max_col, max_row)


def _require_xlsx_extent(sheet: Any, bounds: tuple[int, int, int, int], settings: Settings) -> None:
    _min_col, _min_row, max_col, max_row = bounds
    if max(max_col, sheet.max_column) * max(max_row, sheet.max_row) > settings.max_structured_elements:
        raise StructuredFileError("XLSX edit would exceed max_structured_elements")


def _inspect_xlsx(data: bytes, settings: Settings, range_ref: str | None = None) -> dict[str, Any]:
    _, load_workbook, *_ = _xlsx_modules()
    with warnings.catch_warnings():
        warnings.filterwarnings(
            "ignore",
            message="Unknown extension is not supported and will be removed",
            category=UserWarning,
        )
        book = load_workbook(io.BytesIO(data), read_only=True, data_only=False, keep_links=True)
        try:
            sheets = []
            total = 0
            for sheet in book.worksheets:
                cells = _sheet_cell_count(sheet)
                total += cells
                if total > settings.max_structured_elements:
                    raise StructuredFileError("XLSX cells exceed max_structured_elements")
                preview_range = (
                    range_ref
                    if range_ref and sheet.title == book.active.title
                    else f"A1:{sheet.cell(min(sheet.max_row, 20), min(sheet.max_column, 20)).coordinate}"
                )
                try:
                    preview_range, _ = _xlsx_range_bounds(
                        preview_range, settings, "preview range"
                    )
                    rows = [[cell.value for cell in row] for row in sheet[preview_range]]
                except ValueError as error:
                    raise StructuredFileError("invalid XLSX range") from error
                sheets.append(
                    {
                        "name": sheet.title,
                        "state": sheet.sheet_state,
                        "max_row": sheet.max_row,
                        "max_column": sheet.max_column,
                        "preview_range": preview_range,
                        "values": rows,
                    }
                )
            unsupported = _xlsx_unsupported(data, settings)
            return {
                "format": "xlsx",
                "sheets": sheets,
                "write_rejected_features": unsupported,
                "package_patch_supported_operations": (
                    []
                    if "digital signatures" in unsupported
                    else ["cell_set", "range_set", "range_clear"]
                ),
            }
        finally:
            book.close()


def _xlsx_cell_format(cell: Any, spec: dict[str, Any], styles: Any) -> None:
    Alignment, Border, _Font, PatternFill, Side = styles
    if "font" in spec:
        values = spec["font"]
        if not isinstance(values, dict):
            raise StructuredFileError("font must be an object")
        allowed = {"name", "size", "bold", "italic", "underline", "color"}
        if set(values) - allowed:
            raise StructuredFileError("unsupported font option")
        font = copy(cell.font)
        for key, value in values.items():
            setattr(font, key, value)
        cell.font = font
    if "fill" in spec:
        color = _text(spec["fill"], "fill")
        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            raise StructuredFileError("fill must be six hexadecimal digits")
        cell.fill = PatternFill("solid", fgColor=color.upper())
    if "alignment" in spec:
        values = spec["alignment"]
        if not isinstance(values, dict):
            raise StructuredFileError("alignment must be an object")
        cell.alignment = Alignment(**values)
    if "number_format" in spec:
        cell.number_format = _text(spec["number_format"], "number_format")
    if "border" in spec:
        color = _text(spec["border"], "border")
        side = Side(style="thin", color=color)
        cell.border = Border(left=side, right=side, top=side, bottom=side)


def _xlsx_selected_rows(sheet: Any, range_ref: str) -> tuple[tuple[Any, ...], ...]:
    selected = sheet[range_ref]
    if isinstance(selected, tuple):
        if not selected or isinstance(selected[0], tuple):
            return selected
        return (selected,)
    return ((selected,),)


def _transform_xlsx(
    data: bytes, operations: list[Any], settings: Settings
) -> tuple[bytes, dict[str, Any]]:
    unsupported = _xlsx_unsupported(data, settings)
    if unsupported:
        return _transform_xlsx_package_preserving(data, operations, settings, unsupported)
    Workbook, load_workbook, BarChart, LineChart, Reference, CellIsRule, styles, DataValidation, _ = _xlsx_modules()
    del Workbook
    book = load_workbook(io.BytesIO(data), data_only=False, keep_links=True)
    try:
        if sum(_sheet_cell_count(sheet) for sheet in book.worksheets) > settings.max_structured_elements:
            raise StructuredFileError("XLSX cells exceed max_structured_elements")
        for raw in operations:
            op = _operation(raw)
            name = op["op"]
            if name == "sheet_add":
                title = _text(op.get("title"), "title")
                book.create_sheet(title)
            elif name == "sheet_remove":
                sheet = book[_text(op.get("sheet"), "sheet")]
                if len(book.worksheets) == 1:
                    raise StructuredFileError("cannot remove the last worksheet")
                book.remove(sheet)
            elif name == "sheet_rename":
                book[_text(op.get("sheet"), "sheet")].title = _text(op.get("title"), "title")
            elif name == "sheet_copy":
                copied = book.copy_worksheet(book[_text(op.get("sheet"), "sheet")])
                if "title" in op:
                    copied.title = _text(op["title"], "title")
            elif name == "sheet_move":
                sheet = book[_text(op.get("sheet"), "sheet")]
                position = _index(op.get("position"), "position")
                if position >= len(book.worksheets):
                    raise StructuredFileError("sheet position is outside the workbook")
                current = book.index(sheet)
                book.move_sheet(sheet, offset=position - current)
            else:
                sheet = book[_text(op.get("sheet"), "sheet")]
                if name == "cell_set":
                    cell_ref, bounds = _xlsx_range_bounds(op.get("cell"), settings, "cell")
                    if bounds[0] != bounds[2] or bounds[1] != bounds[3]:
                        raise StructuredFileError("cell must identify exactly one XLSX cell")
                    _require_xlsx_extent(sheet, bounds, settings)
                    cell = sheet[cell_ref]
                    cell.value = op.get("value")
                    if "format" in op:
                        _xlsx_cell_format(cell, op["format"], styles)
                elif name == "range_set":
                    range_ref, bounds = _xlsx_range_bounds(op.get("range"), settings)
                    _require_xlsx_extent(sheet, bounds, settings)
                    values = op.get("values")
                    if not isinstance(values, list) or not all(isinstance(row, list) for row in values):
                        raise StructuredFileError("values must be a two-dimensional array")
                    rows = _xlsx_selected_rows(sheet, range_ref)
                    if len(values) != len(rows) or any(len(value_row) != len(cells) for value_row, cells in zip(values, rows, strict=False)):
                        raise StructuredFileError("values shape must match range")
                    for value_row, cells in zip(values, rows, strict=False):
                        for value, cell in zip(value_row, cells, strict=False):
                            cell.value = value
                elif name == "range_clear":
                    range_ref, bounds = _xlsx_range_bounds(op.get("range"), settings)
                    _require_xlsx_extent(sheet, bounds, settings)
                    for row in _xlsx_selected_rows(sheet, range_ref):
                        for cell in row:
                            cell.value = None
                elif name in {"range_copy", "range_fill"}:
                    from openpyxl.formula.translate import Translator
                    _, source_bounds = _xlsx_range_bounds(op.get("source"), settings, "source")
                    _, target_bounds = _xlsx_range_bounds(op.get("target"), settings, "target")
                    source_min_col, source_min_row, source_max_col, source_max_row = source_bounds
                    target_min_col, target_min_row, target_max_col, target_max_row = target_bounds
                    _require_xlsx_extent(sheet, target_bounds, settings)
                    source_rows = source_max_row - source_min_row + 1
                    source_cols = source_max_col - source_min_col + 1
                    target_rows = target_max_row - target_min_row + 1
                    target_cols = target_max_col - target_min_col + 1
                    if name == "range_copy" and (source_rows, source_cols) != (target_rows, target_cols):
                        raise StructuredFileError("range_copy source and target shapes must match")
                    copy_format = bool(op.get("copy_format", True))
                    source_snapshot = []
                    for row_offset in range(source_rows):
                        snapshot_row = []
                        for column_offset in range(source_cols):
                            source_cell = sheet.cell(
                                source_min_row + row_offset,
                                source_min_col + column_offset,
                            )
                            snapshot_row.append(
                                {
                                    "coordinate": source_cell.coordinate,
                                    "value": source_cell.value,
                                    "style": copy(source_cell._style),
                                    "number_format": source_cell.number_format,
                                    "protection": copy(source_cell.protection),
                                }
                            )
                        source_snapshot.append(snapshot_row)
                    for row_offset in range(target_rows):
                        for column_offset in range(target_cols):
                            source_cell = source_snapshot[row_offset % source_rows][
                                column_offset % source_cols
                            ]
                            target_cell = sheet.cell(target_min_row + row_offset, target_min_col + column_offset)
                            value = source_cell["value"]
                            if isinstance(value, str) and value.startswith("="):
                                value = Translator(
                                    value, origin=source_cell["coordinate"]
                                ).translate_formula(target_cell.coordinate)
                            target_cell.value = value
                            if copy_format:
                                target_cell._style = copy(source_cell["style"])
                                target_cell.number_format = source_cell["number_format"]
                                target_cell.protection = copy(source_cell["protection"])
                elif name in {"rows_insert", "rows_delete", "columns_insert", "columns_delete"}:
                    index = _index(op.get("index"), "index", minimum=1)
                    amount = _index(op.get("amount", 1), "amount", minimum=1)
                    if amount > settings.max_structured_elements:
                        raise StructuredFileError("row or column amount exceeds max_structured_elements")
                    projected_rows = sheet.max_row + amount if name == "rows_insert" else sheet.max_row
                    projected_columns = sheet.max_column + amount if name == "columns_insert" else sheet.max_column
                    if projected_rows * projected_columns > settings.max_structured_elements:
                        raise StructuredFileError("XLSX edit would exceed max_structured_elements")
                    getattr(sheet, name.replace("rows_", "").replace("columns_", "") + ("_rows" if name.startswith("rows") else "_cols"))(index, amount)
                elif name == "merge":
                    range_ref, bounds = _xlsx_range_bounds(op.get("range"), settings)
                    _require_xlsx_extent(sheet, bounds, settings)
                    sheet.merge_cells(range_ref)
                elif name == "unmerge":
                    range_ref, _ = _xlsx_range_bounds(op.get("range"), settings)
                    sheet.unmerge_cells(range_ref)
                elif name == "format_range":
                    spec = op.get("format")
                    if not isinstance(spec, dict):
                        raise StructuredFileError("format must be an object")
                    range_ref, bounds = _xlsx_range_bounds(op.get("range"), settings)
                    _require_xlsx_extent(sheet, bounds, settings)
                    for row in _xlsx_selected_rows(sheet, range_ref):
                        for cell in row:
                            _xlsx_cell_format(cell, spec, styles)
                elif name == "dimensions_set":
                    if "row" in op:
                        dimension = sheet.row_dimensions[_index(op["row"], "row", minimum=1)]
                    elif "column" in op:
                        dimension = sheet.column_dimensions[_text(op["column"], "column")]
                    else:
                        raise StructuredFileError("dimensions_set needs row or column")
                    if "size" in op:
                        if not isinstance(op["size"], (int, float)) or op["size"] <= 0:
                            raise StructuredFileError("size must be positive")
                        dimension.height = op["size"] if "row" in op else None
                        dimension.width = op["size"] if "column" in op else None
                    if "hidden" in op:
                        if not isinstance(op["hidden"], bool):
                            raise StructuredFileError("hidden must be boolean")
                        dimension.hidden = op["hidden"]
                elif name == "freeze_panes_set":
                    cell_ref, bounds = _xlsx_range_bounds(op.get("cell"), settings, "cell")
                    if bounds[0] != bounds[2] or bounds[1] != bounds[3]:
                        raise StructuredFileError("freeze pane must identify one cell")
                    sheet.freeze_panes = cell_ref
                elif name == "autofilter_set":
                    range_ref, _ = _xlsx_range_bounds(op.get("range"), settings)
                    sheet.auto_filter.ref = range_ref
                elif name == "table_add":
                    from openpyxl.worksheet.table import Table, TableStyleInfo
                    range_ref, bounds = _xlsx_range_bounds(op.get("range"), settings)
                    _require_xlsx_extent(sheet, bounds, settings)
                    table = Table(displayName=_text(op.get("name"), "name"), ref=range_ref)
                    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
                    sheet.add_table(table)
                elif name == "validation_add":
                    validation = DataValidation(type=_text(op.get("type", "list"), "type"), formula1=_text(op.get("formula1"), "formula1"), allow_blank=bool(op.get("allow_blank", True)))
                    sheet.add_data_validation(validation)
                    range_ref, _ = _xlsx_range_bounds(op.get("range"), settings)
                    validation.add(range_ref)
                elif name == "conditional_cell_is":
                    color = _text(op.get("fill"), "fill")
                    rule = CellIsRule(operator=_text(op.get("operator"), "operator"), formula=[_text(op.get("formula"), "formula")], fill=styles[3]("solid", fgColor=color))
                    range_ref, _ = _xlsx_range_bounds(op.get("range"), settings)
                    sheet.conditional_formatting.add(range_ref, rule)
                elif name == "chart_add":
                    chart_type = _text(op.get("type", "bar"), "type")
                    if chart_type not in {"bar", "line"}:
                        raise StructuredFileError("chart type must be bar or line")
                    chart = BarChart() if chart_type == "bar" else LineChart()
                    _data_range, bounds = _xlsx_range_bounds(
                        op.get("data_range"), settings, "data range"
                    )
                    min_col, min_row, max_col, max_row = bounds
                    data_ref = Reference(
                        sheet,
                        min_col=min_col,
                        min_row=min_row,
                        max_col=max_col,
                        max_row=max_row,
                    )
                    chart.add_data(data_ref, titles_from_data=bool(op.get("titles_from_data", True)))
                    anchor, anchor_bounds = _xlsx_range_bounds(
                        op.get("anchor", "E2"), settings, "anchor"
                    )
                    if anchor_bounds[0] != anchor_bounds[2] or anchor_bounds[1] != anchor_bounds[3]:
                        raise StructuredFileError("chart anchor must identify one cell")
                    sheet.add_chart(chart, anchor)
                elif name == "page_setup_set":
                    values = op.get("values")
                    if not isinstance(values, dict):
                        raise StructuredFileError("page setup values must be an object")
                    for key in ("orientation", "paperSize", "fitToWidth", "fitToHeight"):
                        if key in values:
                            setattr(sheet.page_setup, key, values[key])
                else:
                    raise StructuredFileError(f"unsupported XLSX operation: {name}")
            if sum(_sheet_cell_count(item) for item in book.worksheets) > settings.max_structured_elements:
                raise StructuredFileError("XLSX cells exceed max_structured_elements")
        output = io.BytesIO()
        book.save(output)
        return output.getvalue(), {"preservation_mode": "library_rewrite"}
    finally:
        book.close()


@dataclass
class CsvDocument:
    rows: list[list[str]]
    encoding: str
    codec: str
    bom: bytes
    dialect: csv.Dialect
    newline: str
    final_newline: bool


def _parse_csv(data: bytes, kind: str, settings: Settings) -> CsvDocument:
    if len(data) > settings.max_structured_file_bytes:
        _require_size(data, settings)
    if data.startswith(b"\xef\xbb\xbf"):
        encoding, codec, bom = "utf-8-sig", "utf-8", b"\xef\xbb\xbf"
    elif data.startswith(b"\xff\xfe"):
        encoding, codec, bom = "utf-16-le", "utf-16-le", b"\xff\xfe"
    elif data.startswith(b"\xfe\xff"):
        encoding, codec, bom = "utf-16-be", "utf-16-be", b"\xfe\xff"
    else:
        encoding, codec, bom = "utf-8", "utf-8", b""
    try:
        text = data[len(bom) :].decode(codec)
    except UnicodeDecodeError as error:
        raise StructuredFileError("CSV/TSV encoding is unsupported or invalid") from error
    without_crlf = text.replace("\r\n", "")
    if "\r" in without_crlf or ("\r\n" in text and "\n" in without_crlf):
        raise StructuredFileError("CSV/TSV contains mixed or ambiguous newline sequences")
    sample = text[:8192]
    if kind == "tsv":
        if "\t" not in sample and any(marker in sample for marker in (",", ";", "|")):
            raise StructuredFileError("TSV delimiter identity is ambiguous")
        dialect = csv.excel_tab
    else:
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error as error:
            if any(marker in sample for marker in (",", ";", "\t", "|")):
                raise StructuredFileError("CSV delimiter identity is ambiguous") from error
            dialect = csv.excel  # A delimiter-free document is an unambiguous single column.
    rows = list(csv.reader(io.StringIO(text, newline=""), dialect))
    _bounded(rows, settings, "CSV rows")
    if sum(len(row) for row in rows) > settings.max_structured_elements:
        raise StructuredFileError("CSV cells exceed max_structured_elements")
    newline = "\r\n" if "\r\n" in text else "\n"
    return CsvDocument(rows, encoding, codec, bom, dialect, newline, text.endswith(("\r\n", "\n")))


def _inspect_csv(data: bytes, kind: str, settings: Settings) -> dict[str, Any]:
    document = _parse_csv(data, kind, settings)
    return {"format": kind, "encoding": document.encoding, "delimiter": document.dialect.delimiter, "quotechar": document.dialect.quotechar, "doublequote": document.dialect.doublequote, "escapechar": document.dialect.escapechar, "newline": document.newline, "final_newline": document.final_newline, "rows": len(document.rows), "columns": max((len(row) for row in document.rows), default=0), "preview": document.rows[:200], "truncated": len(document.rows) > 200}


def _require_csv_bounds(document: CsvDocument, settings: Settings) -> None:
    _bounded(document.rows, settings, "CSV rows")
    if sum(len(row) for row in document.rows) > settings.max_structured_elements:
        raise StructuredFileError("CSV cells exceed max_structured_elements")


def _transform_csv(data: bytes, kind: str, operations: list[Any], settings: Settings) -> bytes:
    document = _parse_csv(data, kind, settings) if data else CsvDocument([], "utf-8", "utf-8", b"", csv.excel_tab if kind == "tsv" else csv.excel, "\n", True)
    for raw in operations:
        op = _operation(raw)
        name = op["op"]
        if name == "cell_set":
            row = _index(op.get("row"), "row")
            column = _index(op.get("column"), "column")
            if row >= settings.max_structured_elements or column >= settings.max_structured_elements:
                raise StructuredFileError("CSV cell index exceeds max_structured_elements")
            existing_length = len(document.rows[row]) if row < len(document.rows) else 0
            current_cells = sum(len(item) for item in document.rows)
            if current_cells - existing_length + max(existing_length, column + 1) > settings.max_structured_elements:
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            while len(document.rows) <= row:
                document.rows.append([])
            while len(document.rows[row]) <= column:
                document.rows[row].append("")
            document.rows[row][column] = _text(op.get("value", ""), "value")
        elif name == "row_append":
            values = op.get("values")
            if not isinstance(values, list):
                raise StructuredFileError("values must be an array")
            if len(values) > settings.max_structured_elements:
                raise StructuredFileError("CSV row exceeds max_structured_elements")
            if sum(len(row) for row in document.rows) + len(values) > settings.max_structured_elements:
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            document.rows.append([_text(value, "value") for value in values])
        elif name == "row_insert":
            row = _index(op.get("row"), "row")
            if row > len(document.rows):
                raise StructuredFileError("row_insert index is outside the document")
            values = op.get("values", [])
            if not isinstance(values, list):
                raise StructuredFileError("values must be an array")
            if len(values) > settings.max_structured_elements:
                raise StructuredFileError("CSV row exceeds max_structured_elements")
            if sum(len(item) for item in document.rows) + len(values) > settings.max_structured_elements:
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            document.rows.insert(row, [_text(value, "value") for value in values])
        elif name == "row_set":
            row = _index(op.get("row"), "row")
            values = op.get("values")
            if not isinstance(values, list):
                raise StructuredFileError("values must be an array")
            if len(values) > settings.max_structured_elements:
                raise StructuredFileError("CSV row exceeds max_structured_elements")
            if (
                sum(len(item) for item in document.rows)
                - len(document.rows[row])
                + len(values)
                > settings.max_structured_elements
            ):
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            document.rows[row] = [_text(value, "value") for value in values]
        elif name == "row_delete":
            del document.rows[_index(op.get("row"), "row")]
        elif name == "column_insert":
            column = _index(op.get("column"), "column")
            if column >= settings.max_structured_elements:
                raise StructuredFileError("CSV column index exceeds max_structured_elements")
            if sum(len(row) for row in document.rows) + len(document.rows) > settings.max_structured_elements:
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            value = _text(op.get("default", ""), "default")
            for row in document.rows:
                row.insert(min(column, len(row)), value)
        elif name == "column_delete":
            column = _index(op.get("column"), "column")
            if column >= settings.max_structured_elements:
                raise StructuredFileError("CSV column index exceeds max_structured_elements")
            for row in document.rows:
                if column < len(row):
                    del row[column]
        elif name == "column_append":
            values = op.get("values", [])
            if not isinstance(values, list) or len(values) > len(document.rows):
                raise StructuredFileError("column values must fit the existing rows")
            if sum(len(row) for row in document.rows) + len(document.rows) > settings.max_structured_elements:
                raise StructuredFileError("CSV edit would exceed max_structured_elements")
            for index, row in enumerate(document.rows):
                row.append(_text(values[index], "value") if index < len(values) else "")
        else:
            raise StructuredFileError(f"unsupported CSV operation: {name}")
        _require_csv_bounds(document, settings)
    stream = io.StringIO(newline="")
    writer = csv.writer(stream, document.dialect, lineterminator=document.newline)
    writer.writerows(document.rows)
    rendered = stream.getvalue()
    if not document.final_newline and rendered.endswith(document.newline):
        rendered = rendered[: -len(document.newline)]
    output = document.bom + rendered.encode(document.codec)
    _require_size(output, settings)
    return output


def _safe_zip_name(name: str) -> str:
    candidate = name.replace("\\", "/")
    path = PureWindowsPath(candidate)
    if not candidate or candidate.startswith(("/", "\\")) or path.is_absolute() or ":" in candidate:
        raise StructuredFileError("ZIP entry path must be relative and not use ADS")
    parts = [part for part in candidate.split("/") if part]
    if not parts or any(part in {".", ".."} or part.endswith((" ", ".")) for part in parts):
        raise StructuredFileError("unsafe ZIP entry path")
    if any(part.split(".", 1)[0].upper() in _ZIP_RESERVED for part in parts):
        raise StructuredFileError("ZIP entry uses a Windows reserved device name")
    return "/".join(parts)


def _validate_zip_paths(paths: list[tuple[str, bool]]) -> None:
    seen: dict[str, bool] = {}
    files: set[str] = set()
    normalized: list[tuple[str, bool]] = []
    for name, is_directory in paths:
        safe = _safe_zip_name(name)
        folded = safe.casefold()
        if folded in seen and seen[folded] != is_directory:
            raise StructuredFileError("ZIP contains a file/directory path collision")
        if folded in seen:
            raise StructuredFileError("ZIP contains duplicate or colliding entry names")
        seen[folded] = is_directory
        normalized.append((safe, is_directory))
        if not is_directory:
            files.add(folded)
    for name, _ in normalized:
        parts = name.casefold().split("/")
        for end in range(1, len(parts)):
            if "/".join(parts[:end]) in files:
                raise StructuredFileError("ZIP contains a file/directory path collision")


def _zip_entries(data: bytes, settings: Settings) -> tuple[zipfile.ZipFile, list[zipfile.ZipInfo]]:
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as error:
        raise StructuredFileError("invalid ZIP archive") from error
    entries = archive.infolist()
    if len(entries) > settings.max_zip_entries:
        archive.close()
        raise StructuredFileError("ZIP entry count exceeds max_zip_entries")
    try:
        _validate_zip_paths([(info.filename, info.is_dir()) for info in entries])
    except Exception:
        archive.close()
        raise
    expanded = sum(info.file_size for info in entries)
    if expanded > settings.max_zip_expanded_bytes:
        archive.close()
        raise StructuredFileError("ZIP expanded size exceeds max_zip_expanded_bytes")
    if any(info.file_size > settings.max_zip_expanded_bytes for info in entries):
        archive.close()
        raise StructuredFileError("ZIP entry expanded size exceeds max_zip_expanded_bytes")
    if any(info.flag_bits & 0x1 for info in entries):
        archive.close()
        raise StructuredFileError("encrypted ZIP entries are unsupported")
    return archive, entries


def _inspect_zip(data: bytes, settings: Settings) -> dict[str, Any]:
    archive, entries = _zip_entries(data, settings)
    try:
        return {"format": "zip", "entry_count": len(entries), "expanded_bytes": sum(info.file_size for info in entries), "entries": [{"name": _safe_zip_name(info.filename), "compressed_bytes": info.compress_size, "bytes": info.file_size, "is_directory": info.is_dir()} for info in entries[:500]], "truncated": len(entries) > 500}
    finally:
        archive.close()


def read_zip_entry(data: bytes, name: str, settings: Settings) -> bytes:
    """Return one prevalidated entry without writing it anywhere."""
    safe_name = _safe_zip_name(name)
    archive, entries = _zip_entries(data, settings)
    try:
        matches = [info for info in entries if not info.is_dir() and _safe_zip_name(info.filename).casefold() == safe_name.casefold()]
        if len(matches) != 1:
            raise StructuredFileError("ZIP entry was not found")
        payload = archive.read(matches[0])
        if len(payload) != matches[0].file_size:
            raise StructuredFileError("ZIP entry size verification failed")
        return payload
    finally:
        archive.close()


def read_zip_entries(
    data: bytes, names: list[str] | None, settings: Settings
) -> dict[str, bytes]:
    """Return a fully validated, bounded set of files for transactional extraction."""
    archive, entries = _zip_entries(data, settings)
    try:
        requested = None if names is None else {_safe_zip_name(name).casefold() for name in names}
        result: dict[str, bytes] = {}
        for info in entries:
            if info.is_dir():
                continue
            safe = _safe_zip_name(info.filename)
            if requested is not None and safe.casefold() not in requested:
                continue
            payload = archive.read(info)
            if len(payload) != info.file_size:
                raise StructuredFileError("ZIP entry size verification failed")
            result[safe] = payload
        if requested is not None and {name.casefold() for name in result} != requested:
            raise StructuredFileError("one or more ZIP entries were not found")
        return result
    finally:
        archive.close()


def _zip_payload(op: dict[str, Any]) -> bytes:
    if "text" in op:
        return _text(op["text"], "text").encode("utf-8")
    if "base64" in op:
        try:
            return base64.b64decode(_text(op["base64"], "base64"), validate=True)
        except ValueError as error:
            raise StructuredFileError("base64 must be valid") from error
    raise StructuredFileError("ZIP entry operation needs text or base64")


def _matching_zip_key(entries: dict[str, tuple[zipfile.ZipInfo | None, bytes]], name: str) -> str | None:
    folded = name.casefold()
    for existing in entries:
        if existing.casefold() == folded:
            return existing
    return None


def _transform_zip(data: bytes, operations: list[Any], settings: Settings) -> bytes:
    entries: dict[str, tuple[zipfile.ZipInfo | None, bytes]] = {}
    directories: list[zipfile.ZipInfo] = []
    if data:
        archive, infos = _zip_entries(data, settings)
        try:
            for info in infos:
                if info.is_dir():
                    directories.append(copy(info))
                else:
                    entries[_safe_zip_name(info.filename)] = (info, archive.read(info))
        finally:
            archive.close()
    for raw in operations:
        op = _operation(raw)
        name = op["op"]
        if name in {"entry_add", "entry_replace"}:
            entry = _safe_zip_name(_text(op.get("name"), "name"))
            existing = _matching_zip_key(entries, entry)
            if name == "entry_add":
                if existing is not None:
                    raise StructuredFileError("ZIP entry already exists")
                entries[entry] = (None, _zip_payload(op))
            else:
                if existing is None:
                    raise StructuredFileError("ZIP entry does not exist")
                entries[existing] = (None, _zip_payload(op))
        elif name == "entry_delete":
            entry = _safe_zip_name(_text(op.get("name"), "name"))
            existing = _matching_zip_key(entries, entry)
            if existing is None:
                raise StructuredFileError("ZIP entry does not exist")
            del entries[existing]
        else:
            raise StructuredFileError(f"unsupported ZIP operation: {name}")
    _validate_zip_paths(
        [(info.filename, True) for info in directories] + [(name, False) for name in entries]
    )
    if len(entries) + len(directories) > settings.max_zip_entries or sum(len(value[1]) for value in entries.values()) > settings.max_zip_expanded_bytes:
        raise StructuredFileError("resulting ZIP exceeds configured limits")
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, allowZip64=False) as archive:
        for info in directories:
            archive.writestr(copy(info), b"")
        for name, (info, payload) in sorted(entries.items()):
            if info is not None:
                clone = copy(info)
                clone.filename = name
                archive.writestr(clone, payload)
            else:
                archive.writestr(name, payload)
    result = output.getvalue()
    verified, _ = _zip_entries(result, settings)
    verified.close()
    return result


def _inspect_image(data: bytes, settings: Settings) -> dict[str, Any]:
    Image, UnidentifiedImageError = _image_module()
    Image.MAX_IMAGE_PIXELS = settings.max_image_pixels
    try:
        with Image.open(io.BytesIO(data)) as image:
            image.verify()
        with Image.open(io.BytesIO(data)) as image:
            pixels = image.width * image.height
            if pixels > settings.max_image_pixels:
                raise StructuredFileError("image pixel count exceeds max_image_pixels")
            if pixels * 8 > settings.max_image_decoded_bytes:
                raise StructuredFileError("decoded image exceeds max_image_decoded_bytes")
            metadata = {str(key): str(value)[:500] for key, value in image.info.items() if key.casefold() not in {"exif", "icc_profile"}}
            return {"format": "image", "image_format": image.format, "width": image.width, "height": image.height, "mode": image.mode, "frame_count": getattr(image, "n_frames", 1), "metadata": metadata, "metadata_policy": {"default": "preserve EXIF, ICC, and DPI when the destination encoder supports them", "metadata_remove": "remove EXIF; ICC and DPI are also omitted"}}
    except UnidentifiedImageError as error:
        raise StructuredFileError("unsupported or corrupt image") from error


def _require_image_pixels(width: int, height: int, settings: Settings) -> None:
    if width <= 0 or height <= 0 or width * height > settings.max_image_pixels:
        raise StructuredFileError("image pixel count exceeds max_image_pixels")
    if width * height * 8 > settings.max_image_decoded_bytes:
        raise StructuredFileError("decoded image exceeds max_image_decoded_bytes")


def _transform_image(data: bytes, operations: list[Any], settings: Settings) -> tuple[bytes, str | None]:
    Image, UnidentifiedImageError = _image_module()
    Image.MAX_IMAGE_PIXELS = settings.max_image_pixels
    try:
        image = Image.open(io.BytesIO(data))
    except UnidentifiedImageError as error:
        raise StructuredFileError("unsupported or corrupt image") from error
    try:
        _require_image_pixels(image.width, image.height, settings)
        if getattr(image, "n_frames", 1) != 1:
            raise StructuredFileError(
                "multi-frame image transformation is unsupported; use the container processing path"
            )
        source_format = image.format
        source_exif = image.info.get("exif")
        source_icc = image.info.get("icc_profile")
        source_dpi = image.info.get("dpi")
        output_format: str | None = None
        quality: int | None = None
        strip_metadata = False
        for raw in operations:
            op = _operation(raw)
            name = op["op"]
            if name == "resize":
                width, height = _index(op.get("width"), "width", minimum=1), _index(op.get("height"), "height", minimum=1)
                _require_image_pixels(width, height, settings)
                image = image.resize((width, height))
            elif name == "thumbnail":
                width, height = _index(op.get("max_width"), "max_width", minimum=1), _index(op.get("max_height"), "max_height", minimum=1)
                image.thumbnail((width, height))
            elif name == "crop":
                box = op.get("box")
                if not isinstance(box, list) or len(box) != 4:
                    raise StructuredFileError("crop box must be [left, top, right, bottom]")
                left, top, right, bottom = tuple(_index(value, "crop coordinate") for value in box)
                if not (left < right <= image.width and top < bottom <= image.height):
                    raise StructuredFileError("crop box must stay within the source image")
                image = image.crop((left, top, right, bottom))
            elif name == "rotate":
                angle = op.get("degrees")
                if not isinstance(angle, (int, float)):
                    raise StructuredFileError("degrees must be a number")
                image = image.rotate(angle, expand=bool(op.get("expand", True)))
            elif name == "flip":
                direction = _text(op.get("direction"), "direction")
                if direction == "horizontal":
                    image = image.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
                elif direction == "vertical":
                    image = image.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
                else:
                    raise StructuredFileError("flip direction must be horizontal or vertical")
            elif name == "convert":
                output_format = _text(op.get("format"), "format").upper()
                if output_format not in {"PNG", "JPEG", "WEBP", "GIF", "BMP", "TIFF"}:
                    raise StructuredFileError("unsupported image output format")
            elif name == "quality":
                quality = _index(op.get("value"), "quality", minimum=1)
                if quality > 100:
                    raise StructuredFileError("quality must be <= 100")
            elif name == "metadata_remove":
                strip_metadata = True
            else:
                raise StructuredFileError(f"unsupported image operation: {name}")
            _require_image_pixels(image.width, image.height, settings)
        output_format = output_format or source_format or "PNG"
        if output_format == "JPEG" and image.mode not in {"RGB", "L"}:
            image = image.convert("RGB")
        output = io.BytesIO()
        kwargs: dict[str, Any] = {}
        if quality is not None and output_format in {"JPEG", "WEBP"}:
            kwargs["quality"] = quality
        if strip_metadata:
            kwargs["exif"] = b""
        else:
            if isinstance(source_exif, bytes):
                kwargs["exif"] = source_exif
            if isinstance(source_icc, bytes):
                kwargs["icc_profile"] = source_icc
            if isinstance(source_dpi, tuple) and len(source_dpi) == 2:
                kwargs["dpi"] = source_dpi
        image.save(output, format=output_format, **kwargs)
        return output.getvalue(), output_format
    finally:
        image.close()


def inspect(data: bytes, path: str, settings: Settings, *, format: str | None = None, range_ref: str | None = None) -> dict[str, Any]:
    _require_size(data, settings)
    kind = infer_format(path, format)
    if kind == "docx": result = _inspect_docx(data, settings)
    elif kind == "xlsx": result = _inspect_xlsx(data, settings, range_ref)
    elif kind in {"csv", "tsv"}: result = _inspect_csv(data, kind, settings)
    elif kind == "zip": result = _inspect_zip(data, settings)
    else: result = _inspect_image(data, settings)
    result.update(bytes=len(data), sha256=sha256_bytes(data))
    return result


def transform(data: bytes, path: str, operations: list[Any], settings: Settings, *, format: str | None = None) -> tuple[bytes, dict[str, Any]]:
    if not isinstance(operations, list) or not operations:
        raise StructuredFileError("operations must be a non-empty array")
    if len(operations) > min(settings.max_structured_elements, 1000):
        raise StructuredFileError("operations exceed the bounded processing limit")
    _require_size(data, settings)
    kind = infer_format(path, format)
    if kind == "docx": output, extra = _transform_docx(data, operations, settings)
    elif kind == "xlsx": output, extra = _transform_xlsx(data, operations, settings)
    elif kind in {"csv", "tsv"}: output, extra = _transform_csv(data, kind, operations, settings), {}
    elif kind == "zip": output, extra = _transform_zip(data, operations, settings), {}
    else:
        output, image_format = _transform_image(data, operations, settings)
        format_extensions = {
            "PNG": {".png"}, "JPEG": {".jpg", ".jpeg"}, "WEBP": {".webp"},
            "GIF": {".gif"}, "BMP": {".bmp"}, "TIFF": {".tif", ".tiff"},
        }
        if image_format is not None and PureWindowsPath(path).suffix.casefold() not in format_extensions[image_format]:
            raise StructuredFileError(
                "image conversion would not match the target extension; use a matching output path"
            )
        extra = {"image_format": image_format}
    _require_size(output, settings)
    return output, {"format": kind, "operations": [str(_operation(item)["op"]) for item in operations], **extra}
