from __future__ import annotations

import base64
import importlib
import io
import re
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from PIL import Image

from windows_local_mcp.util import sha256_bytes


def load_server(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    root = tmp_path / "workspace"
    root.mkdir()
    data = tmp_path / "data"
    config = tmp_path / "config.toml"
    config.write_text(
        "\n".join(
            [
                f'workspace_root = "{str(root).replace(chr(92), chr(92) * 2)}"',
                f'data_dir = "{str(data).replace(chr(92), chr(92) * 2)}"',
                "protect_data_dir_acl = false",
                "git_enabled = false",
                "max_structured_file_bytes = 1048576",
                "max_transfer_chunk_bytes = 4096",
            ]
        ),
        encoding="utf-8",
    )
    monkeypatch.setenv("LOCAL_MCP_CONFIG", str(config))
    monkeypatch.delenv("LOCAL_MCP_ROOT", raising=False)
    sys.modules.pop("windows_local_mcp.server", None)
    return importlib.import_module("windows_local_mcp.server"), root


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Heading", 1)
    document.add_paragraph("original text")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def xlsx_bytes() -> bytes:
    book = Workbook()
    sheet = book.active
    sheet.title = "Data"
    sheet.append(["Name", "Amount"])
    sheet.append(["A", 2])
    output = io.BytesIO()
    book.save(output)
    return output.getvalue()


def remove_xlsx_dimensions(data: bytes) -> bytes:
    """Match producers that omit the optional worksheet dimension metadata."""
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as source, zipfile.ZipFile(output, "w") as target:
        target.comment = source.comment
        for info in source.infolist():
            payload = source.read(info)
            if info.filename.startswith("xl/worksheets/"):
                payload = re.sub(br"<dimension\b[^>]*/>", b"", payload, count=1)
            target.writestr(info, payload)
    return output.getvalue()


def add_package_part(data: bytes, name: str, payload: bytes) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as source, zipfile.ZipFile(output, "w") as target:
        target.comment = source.comment
        for info in source.infolist():
            target.writestr(info, source.read(info))
        target.writestr(name, payload)
    return output.getvalue()


def replace_package_part(data: bytes, name: str, update) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as source, zipfile.ZipFile(output, "w") as target:
        target.comment = source.comment
        for info in source.infolist():
            payload = source.read(info)
            target.writestr(info, update(payload) if info.filename == name else payload)
    return output.getvalue()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_docx_xlsx_and_csv_use_the_checkpointed_local_path(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    (root / "report.docx").write_bytes(docx_bytes())
    before = sha256_bytes((root / "report.docx").read_bytes())
    result = server.structured_file_apply(
        "report.docx",
        [
            {"op": "paragraph_update", "index": 1, "text": "edited text", "format": {"run": {"font_size_pt": 10.5}}},
            {"op": "table_cell_set", "table": 0, "row": 0, "column": 1, "text": "updated"},
            {"op": "metadata_set", "values": {"title": "Updated"}},
        ],
        expected_sha256=before,
        reason="edit report",
    )
    assert result["execution_path"] == "broker_direct"
    assert result["rollback_state"] == "complete"
    record = server.runtime.audit.get_operation(result["operation_id"])
    assert record["tier"] == "structured_processing"
    assert record["child_pid"] is None
    document = Document(root / "report.docx")
    assert document.paragraphs[1].text == "edited text"
    assert document.tables[0].cell(0, 1).text == "updated"
    assert document.core_properties.title == "Updated"

    (root / "table.xlsx").write_bytes(xlsx_bytes())
    result = server.structured_file_apply(
        "table.xlsx",
        [
            {"op": "cell_set", "sheet": "Data", "cell": "C2", "value": "=SUM(B2:B2)"},
            {"op": "format_range", "sheet": "Data", "range": "A1:C1", "format": {"font": {"bold": True}, "fill": "DDEEFF"}},
            {"op": "freeze_panes_set", "sheet": "Data", "cell": "A2"},
            {"op": "autofilter_set", "sheet": "Data", "range": "A1:C2"},
        ],
        expected_sha256=sha256_bytes((root / "table.xlsx").read_bytes()),
    )
    assert result["format"] == "xlsx"
    book = load_workbook(root / "table.xlsx", data_only=False)
    assert book["Data"]["C2"].value == "=SUM(B2:B2)"
    assert book["Data"]["A1"].font.bold is True
    assert book["Data"].freeze_panes == "A2"

    created = server.structured_file_apply(
        "new.csv", [{"op": "row_append", "values": ["name", "value"]}, {"op": "row_append", "values": ["A", "1"]}], reason="create csv"
    )
    assert created["before_bytes"] == 0
    inspected = server.structured_file_inspect("new.csv")
    assert inspected["preview"] == [["name", "value"], ["A", "1"]]
    session = server.session_info()
    processing = session["structured_file_processing"]
    assert "chatgpt_container" in processing
    assert "broker_direct" in processing
    assert "external-process use alone does not require Codex Sandbox" in processing[
        "external_processing_policy"
    ]
    assert session["transport"]["stdio"]["available"] is True
    assert session["transport"]["http"]["available"] is False
    assert session["transport"]["http"]["startup_validation"] == "rejected when configured"
    sandbox = session["capabilities"]["status"]["codex_sandbox"]
    assert "available" in sandbox
    assert "execution_route_available" in sandbox
    assert "windows_live_verified" in sandbox
    assert set(sandbox["properties"]) == {
        "filesystem_read",
        "filesystem_write",
        "protected_information_read",
        "internet",
        "lan",
        "loopback",
        "descendant_containment",
        "termination",
        "resource_bound",
    }


def test_xlsx_inspection_sizes_unsized_sheets_and_accepts_qualified_range(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    monkeypatch.setattr(server, "assert_control_plane_healthy", lambda _settings: None)
    book = Workbook()
    summary = book.active
    summary.title = "Summary"
    summary["A1"] = "summary"
    verification = book.create_sheet("検証情報")
    verification["A1"] = "項目"
    verification["B8"] = "完了"
    output = io.BytesIO()
    book.save(output)
    (root / "unsized.xlsx").write_bytes(remove_xlsx_dimensions(output.getvalue()))

    inspected = server.structured_file_inspect("unsized.xlsx")
    assert [(sheet["max_row"], sheet["max_column"]) for sheet in inspected["sheets"]] == [
        (1, 1),
        (8, 2),
    ]

    selected = server.structured_file_inspect(
        "unsized.xlsx", range_ref="検証情報!A1:B8"
    )
    verification_result = next(
        sheet for sheet in selected["sheets"] if sheet["name"] == "検証情報"
    )
    assert verification_result["preview_range"] == "A1:B8"
    assert verification_result["values"][0] == ["項目", None]
    assert verification_result["values"][7] == [None, "完了"]


def test_structured_inspect_releases_verified_handle_before_parser_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    monkeypatch.setattr(server, "assert_control_plane_healthy", lambda _settings: None)
    (root / "table.csv").write_bytes(b"a,b\n1,2\n")
    released: list[Path] = []
    real_release = server.release_verified_hold

    def record_release(path: Path) -> None:
        released.append(path)
        real_release(path)

    monkeypatch.setattr(server, "release_verified_hold", record_release)
    monkeypatch.setattr(
        server,
        "inspect_structured",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("parser failed")),
    )

    with pytest.raises(RuntimeError, match="parser failed"):
        server.structured_file_inspect("table.csv")

    assert [Path(str(path)).name for path in released] == ["table.csv"]
    replacement = b"a,b\n3,4\n"
    upload = server.artifact_upload_begin(
        "table.csv",
        len(replacement),
        sha256_bytes(replacement),
        expected_sha256=sha256_bytes(b"a,b\n1,2\n"),
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(replacement).decode("ascii")
    )
    committed = server.artifact_upload_commit(upload["transfer_id"])
    assert committed["after_sha256"] == sha256_bytes(replacement)
    assert (root / "table.csv").read_bytes() == replacement


def test_zip_image_and_transfer_boundaries(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w") as output:
        output.writestr("notes/a.txt", "hello")
    (root / "bundle.zip").write_bytes(archive.getvalue())
    inspected = server.structured_file_inspect("bundle.zip")
    assert inspected["entries"][0]["name"] == "notes/a.txt"
    server.structured_file_apply(
        "bundle.zip",
        [{"op": "entry_add", "name": "notes/b.txt", "text": "world"}],
        expected_sha256=sha256_bytes((root / "bundle.zip").read_bytes()),
    )
    with zipfile.ZipFile(root / "bundle.zip") as output:
        assert output.read("notes/b.txt") == b"world"
    archive_sha = sha256_bytes((root / "bundle.zip").read_bytes())
    extracted = server.zip_entry_extract(
        "bundle.zip", "notes/b.txt", "extracted.txt", expected_archive_sha256=archive_sha
    )
    assert extracted["operation"] == "entry_extract"
    assert (root / "extracted.txt").read_bytes() == b"world"
    assert base64.b64decode(server.zip_entry_read("bundle.zip", "notes/a.txt")["base64"]) == b"hello"
    with pytest.raises(ValueError, match="unsafe ZIP"):
        server.structured_file_apply(
            "bundle.zip",
            [{"op": "entry_add", "name": "../bad", "text": "x"}],
            expected_sha256=sha256_bytes((root / "bundle.zip").read_bytes()),
        )

    image = Image.new("RGB", (2000, 1000), color="red")
    image.save(root / "photo.jpg", exif=b"Exif\x00\x00test")
    result = server.structured_file_apply(
        "photo.jpg",
        [{"op": "resize", "width": 1200, "height": 600}, {"op": "metadata_remove"}, {"op": "quality", "value": 80}],
        expected_sha256=sha256_bytes((root / "photo.jpg").read_bytes()),
    )
    assert result["format"] == "image"
    assert server.structured_file_inspect("photo.jpg")["width"] == 1200

    (root / "download.csv").write_bytes(b"a,b\r\n1,2\r\n")
    download = server.structured_file_download_begin("download.csv", chunk_bytes=4096)
    chunk = server.structured_file_download_chunk(download["transfer_id"], 0)
    assert base64.b64decode(chunk["base64"]) == b"a,b\r\n1,2\r\n"
    payload = b"a,b\r\n3,4\r\n"
    upload = server.structured_file_upload_begin(
        "download.csv", len(payload), sha256_bytes(payload), expected_sha256=download["sha256"]
    )
    server.structured_file_upload_chunk(upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii"))
    committed = server.structured_file_upload_commit(upload["transfer_id"], reason="return edited artifact")
    assert committed["execution_path"] == "transfer"
    assert (root / "download.csv").read_bytes() == payload


def test_transfer_uses_stable_snapshot_and_rejects_incomplete_uploads(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    (root / "source.csv").write_bytes(b"a\n")
    stale = server.structured_file_download_begin("source.csv")
    (root / "source.csv").write_bytes(b"changed\n")
    snapshot = server.structured_file_download_chunk(stale["transfer_id"], 0)
    assert base64.b64decode(snapshot["base64"]) == b"a\n"
    assert sha256_bytes(base64.b64decode(snapshot["base64"])) == snapshot["sha256"]
    audit = server.runtime.audit.get_operation(stale["operation_id"])
    assert any(event["event_type"] == "artifact_download_chunk" for event in audit["events"])
    assert not any(
        operation["tool_name"] == "artifact_download_chunk"
        for operation in server.runtime.audit.list_operations(limit=100)
    )

    payload = b"complete\n"
    upload = server.structured_file_upload_begin("new.tsv", len(payload), sha256_bytes(payload))
    with pytest.raises(RuntimeError, match="incomplete"):
        server.structured_file_upload_commit(upload["transfer_id"])


def test_download_snapshot_tampering_is_rejected(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    (root / "source.bin").write_bytes(b"trusted snapshot")
    download = server.artifact_download_begin("source.bin")
    snapshot = (
        server.runtime.settings.data_dir
        / "binary-transfers"
        / download["transfer_id"]
        / "payload.bin"
    )
    snapshot.write_bytes(b"tampered bytes!!")

    with pytest.raises(RuntimeError, match="snapshot changed"):
        server.artifact_download_chunk(download["transfer_id"], 0)


def test_generic_artifact_transfer_binds_distinct_source_until_commit(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    source = root / "source.pdf"
    source.write_bytes(b"opaque-pdf-source")
    download = server.artifact_download_begin("source.pdf")
    payload = b"processed-pdf-result"
    upload = server.artifact_upload_begin(
        "result.pdf",
        len(payload),
        sha256_bytes(payload),
        source_transfer_id=download["transfer_id"],
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii")
    )
    source.write_bytes(b"changed-after-container-read")
    with pytest.raises(RuntimeError, match="bound source changed"):
        server.artifact_upload_commit(upload["transfer_id"])
    assert not (root / "result.pdf").exists()


def test_distinct_source_change_during_result_commit_rolls_back_output(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    source = root / "source.pdf"
    source.write_bytes(b"bound-source")
    download = server.artifact_download_begin("source.pdf")
    payload = b"derived-result"
    upload = server.artifact_upload_begin(
        "result.pdf",
        len(payload),
        sha256_bytes(payload),
        source_transfer_id=download["transfer_id"],
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii")
    )
    real_verify = server._verify_binary_source_bindings
    calls = 0

    def change_source_on_postflight(bindings):
        nonlocal calls
        calls += 1
        if calls == 3:
            source.write_bytes(b"third-party-change")
        return real_verify(bindings)

    monkeypatch.setattr(
        server, "_verify_binary_source_bindings", change_source_on_postflight
    )

    with pytest.raises(RuntimeError, match="bound source changed"):
        server.artifact_upload_commit(upload["transfer_id"])

    assert source.read_bytes() == b"third-party-change"
    assert not (root / "result.pdf").exists()


def test_generic_artifact_transfer_commits_unknown_structured_format_as_opaque_bytes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    source = root / "source.pptx"
    source.write_bytes(b"opaque-presentation-source")
    download = server.artifact_download_begin("source.pptx")
    payload = b"processed-presentation-result"
    upload = server.artifact_upload_begin(
        "result.pptx",
        len(payload),
        sha256_bytes(payload),
        source_transfer_id=download["transfer_id"],
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii")
    )

    result = server.artifact_upload_commit(upload["transfer_id"])

    assert (root / "result.pptx").read_bytes() == payload
    assert result["artifact_kind"] == "opaque_binary"
    assert result["embedded_code_executed"] is False


def test_artifact_upload_commits_a_new_file_below_a_subdirectory(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    monkeypatch.setattr(server, "assert_control_plane_healthy", lambda _settings: None)
    (root / "results").mkdir()
    payload = b"structured-upload-result"
    upload = server.artifact_upload_begin(
        "results/bin-10.bin", len(payload), sha256_bytes(payload)
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii")
    )

    result = server.artifact_upload_commit(upload["transfer_id"])

    assert result["path"] == "results/bin-10.bin"
    assert (root / "results" / "bin-10.bin").read_bytes() == payload
    assert server.workspace_recovery_required(server.runtime.settings) is False


def test_recovered_artifact_commit_failure_does_not_latch_workspace_mutations(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    monkeypatch.setattr(server, "assert_control_plane_healthy", lambda _settings: None)
    payload = b"temporary-upload"
    upload = server.artifact_upload_begin("recovered.bin", len(payload), sha256_bytes(payload))
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload).decode("ascii")
    )
    real_capture = server.capture_workspace_state

    def fail_after_replacement(settings, operation_id, stage, *, paths=None):
        if stage == "after":
            raise RuntimeError("forced post-write failure")
        return real_capture(settings, operation_id, stage, paths=paths)

    monkeypatch.setattr(server, "capture_workspace_state", fail_after_replacement)
    with pytest.raises(server.WorkspaceMutationError) as captured:
        server.artifact_upload_commit(upload["transfer_id"])

    assert captured.value.recovery_state == "failed_recovered"
    assert not (root / "recovered.bin").exists()
    assert server.workspace_recovery_required(server.runtime.settings) is False

    monkeypatch.setattr(server, "capture_workspace_state", real_capture)
    server.write_file("next.txt", "mutation remains available")
    assert (root / "next.txt").read_text(encoding="utf-8") == "mutation remains available"


def test_artifact_upload_reserves_data_quota_before_accepting_session(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, _root = load_server(tmp_path, monkeypatch)
    from windows_local_mcp.resources import directory_size

    used = directory_size(server.runtime.settings.data_dir)
    server.runtime.settings.max_data_dir_bytes = used + 8192
    payload = b"x" * 8192

    with pytest.raises(RuntimeError, match="quota exceeded"):
        server.artifact_upload_begin(
            "result.pdf", len(payload), sha256_bytes(payload)
        )


def test_artifact_upload_reserves_payload_once_and_chunks_without_quota_rescan(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, _root = load_server(tmp_path, monkeypatch)
    payload = b"x" * 8192
    upload = server.artifact_upload_begin("result.pdf", len(payload), sha256_bytes(payload))
    transfer_root = server.runtime.settings.data_dir / "binary-transfers" / upload["transfer_id"]
    assert (transfer_root / "payload.bin").stat().st_size == len(payload)

    monkeypatch.setattr(
        server,
        "enforce_data_quota",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("upload chunk rescanned data_dir quota")
        ),
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 0, base64.b64encode(payload[:4096]).decode("ascii")
    )
    server.artifact_upload_chunk(
        upload["transfer_id"], 4096, base64.b64encode(payload[4096:]).decode("ascii")
    )
    audit = server.runtime.audit.get_operation(upload["operation_id"])
    assert sum(
        event["event_type"] == "artifact_upload_chunk" for event in audit["events"]
    ) == 2
    assert not any(
        operation["tool_name"] == "artifact_upload_chunk"
        for operation in server.runtime.audit.list_operations(limit=100)
    )


def test_csv_preserves_bom_delimiter_quotes_crlf_and_final_newline_identity(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    original = b'\xef\xbb\xbf"name";"value"\r\n"a";"1"'
    path = root / "identity.csv"
    path.write_bytes(original)
    inspected = server.structured_file_inspect("identity.csv")
    assert inspected["encoding"] == "utf-8-sig"
    assert inspected["delimiter"] == ";"
    assert inspected["quotechar"] == '"'
    assert inspected["newline"] == "\r\n"
    assert inspected["final_newline"] is False

    server.structured_file_apply(
        "identity.csv",
        [{"op": "cell_set", "row": 1, "column": 1, "value": "2"}],
        expected_sha256=sha256_bytes(original),
    )
    updated = path.read_bytes()
    assert updated.startswith(b"\xef\xbb\xbf")
    assert b";" in updated and b"\r\n" in updated
    assert not updated.endswith(b"\r\n")
    capabilities = inspected["preservation_capabilities"]
    assert capabilities["semantic_cells"] == "preserved_except_declared_edits"
    assert capabilities["lexical_quoting"] == "not_preserved_writer_rewrite"
    assert capabilities["byte_identity"] == "not_preserved_after_edit"


def test_image_conversion_uses_distinct_hash_bound_output_path(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    source = root / "source.png"
    Image.new("RGBA", (12, 8), color=(255, 0, 0, 128)).save(source)
    original = source.read_bytes()
    real_lock = server.WorkspaceExecutionLock
    locked_targets: list[tuple[Path, ...] | None] = []

    class RecordingLock:
        def __init__(self, settings, timeout=30.0, *, target=None, targets=None):
            selected = tuple(targets) if targets is not None else ((target,) if target else None)
            locked_targets.append(selected)
            self._lock = real_lock(settings, timeout, target=target, targets=targets)

        def __enter__(self):
            return self._lock.__enter__()

        def __exit__(self, *args):
            return self._lock.__exit__(*args)

    monkeypatch.setattr(server, "WorkspaceExecutionLock", RecordingLock)

    result = server.structured_file_apply(
        "source.png",
        [{"op": "convert", "format": "JPEG"}],
        expected_sha256=sha256_bytes(original),
        output_path="result.jpg",
    )

    output = root / "result.jpg"
    assert source.read_bytes() == original
    with Image.open(output) as converted:
        assert converted.format == "JPEG"
        assert converted.size == (12, 8)
    assert result["path"] == "result.jpg"
    assert result["checkpoint_scope"] == {
        "kind": "paths",
        "paths": ["result.jpg"],
    }
    assert locked_targets == [
        (source.resolve(),),
        (output.resolve(),),
        (output.resolve(), source.resolve()),
    ]

    with pytest.raises(RuntimeError, match="expected_output_sha256 mismatch"):
        server.structured_file_apply(
            "source.png",
            [{"op": "convert", "format": "JPEG"}],
            expected_sha256=sha256_bytes(original),
            output_path="result.jpg",
            expected_output_sha256="0" * 64,
        )


def test_structured_edits_reject_oversized_ranges_before_materialization(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    workbook = root / "bounded.xlsx"
    workbook.write_bytes(xlsx_bytes())
    with pytest.raises(Exception, match="max_structured_elements"):
        server.structured_file_apply(
            "bounded.xlsx",
            [
                {
                    "op": "range_clear",
                    "sheet": "Data",
                    "range": "A1:XFD1048576",
                }
            ],
            expected_sha256=sha256_bytes(workbook.read_bytes()),
        )

    csv_path = root / "bounded.csv"
    csv_path.write_bytes(b"a,b\n")
    with pytest.raises(Exception, match="max_structured_elements"):
        server.structured_file_apply(
            "bounded.csv",
            [
                {
                    "op": "cell_set",
                    "row": 0,
                    "column": server.runtime.settings.max_structured_elements,
                    "value": "x",
                }
            ],
            expected_sha256=sha256_bytes(csv_path.read_bytes()),
        )


def test_office_packages_apply_expanded_size_bounds_before_parsing(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    server.runtime.settings.max_zip_expanded_bytes = 1024
    (root / "expanded.docx").write_bytes(docx_bytes())
    (root / "expanded.xlsx").write_bytes(xlsx_bytes())

    with pytest.raises(ValueError, match="DOCX package exceeds max_zip_expanded_bytes"):
        server.structured_file_inspect("expanded.docx")
    with pytest.raises(ValueError, match="XLSX package exceeds max_zip_expanded_bytes"):
        server.structured_file_inspect("expanded.xlsx")


def test_zip_multi_extract_is_one_checkpointed_transaction(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    package = io.BytesIO()
    with zipfile.ZipFile(package, "w") as archive:
        archive.writestr("one.txt", b"one")
        archive.writestr("nested/two.txt", b"two")
    source = root / "bundle.zip"
    source.write_bytes(package.getvalue())
    real_lock = server.WorkspaceExecutionLock
    locked_targets: list[tuple[Path, ...] | None] = []

    class RecordingLock:
        def __init__(self, settings, timeout=30.0, *, target=None, targets=None):
            selected = tuple(targets) if targets is not None else ((target,) if target else None)
            locked_targets.append(selected)
            self._lock = real_lock(settings, timeout, target=target, targets=targets)

        def __enter__(self):
            return self._lock.__enter__()

        def __exit__(self, *args):
            return self._lock.__exit__(*args)

    monkeypatch.setattr(server, "WorkspaceExecutionLock", RecordingLock)

    result = server.zip_extract_many(
        "bundle.zip",
        "expanded",
        sha256_bytes(source.read_bytes()),
    )
    assert result["extracted_file_count"] == 2
    assert result["rollback_state"] == "complete"
    assert (root / "expanded" / "one.txt").read_bytes() == b"one"
    assert (root / "expanded" / "nested" / "two.txt").read_bytes() == b"two"
    assert locked_targets == [
        (source.resolve(),),
        (
            source.resolve(),
            root / "expanded" / "one.txt",
            root / "expanded" / "nested" / "two.txt",
        ),
    ]


def test_zip_multi_extract_rolls_back_outputs_when_source_changes_after_apply(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    package = io.BytesIO()
    with zipfile.ZipFile(package, "w") as archive:
        archive.writestr("one.txt", b"one")
        archive.writestr("nested/two.txt", b"two")
    source = root / "bundle.zip"
    original = package.getvalue()
    source.write_bytes(original)
    real_verify = server._verify_binary_source_bindings
    calls = 0

    def change_source_on_postflight(bindings):
        nonlocal calls
        calls += 1
        if calls == 2:
            source.write_bytes(b"third-party-change")
        return real_verify(bindings)

    monkeypatch.setattr(server, "_verify_binary_source_bindings", change_source_on_postflight)

    with pytest.raises(RuntimeError, match="bound source changed"):
        server.zip_extract_many(
            "bundle.zip",
            "expanded",
            sha256_bytes(original),
        )

    assert source.read_bytes() == b"third-party-change"
    assert not (root / "expanded" / "one.txt").exists()
    assert not (root / "expanded" / "nested" / "two.txt").exists()


def test_existing_structured_edit_requires_source_hash(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    (root / "existing.csv").write_bytes(b"a,b\n1,2\n")
    with pytest.raises(ValueError, match="expected_sha256 is required"):
        server.structured_file_apply(
            "existing.csv",
            [{"op": "cell_set", "row": 1, "column": 1, "value": "3"}],
        )
    assert (root / "existing.csv").read_bytes() == b"a,b\n1,2\n"


def test_binary_mutation_rejects_change_during_transform(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    target = root / "race.csv"
    target.write_bytes(b"original\n")
    before = sha256_bytes(target.read_bytes())

    def transform(source: bytes):
        assert source == b"original\n"
        target.write_bytes(b"manual change\n")
        return b"model edit\n", {"format": "csv"}

    with pytest.raises(RuntimeError, match="stale or concurrently modified"):
        server._atomic_binary_mutation(
            tool_name="test_race",
            path="race.csv",
            expected_sha256=before,
            reason="test",
            request_summary={},
            transform=transform,
            require_expected_for_existing=True,
        )
    assert target.read_bytes() == b"manual change\n"


def test_docx_run_formatting_is_scoped_and_section_size_is_inspectable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("first")
    second = paragraph.add_run("second")
    first.bold = False
    second.bold = False
    section = document.sections[0]
    section.header.paragraphs[0].text = "header text"
    output = io.BytesIO()
    document.save(output)
    path = root / "runs.docx"
    path.write_bytes(output.getvalue())

    server.structured_file_apply(
        "runs.docx",
        [
            {"op": "run_update", "paragraph": 0, "run": 0, "format": {"bold": True}},
            {"op": "section_set", "section": 0, "page_width_inches": 8.0, "page_height_inches": 10.0},
        ],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )
    edited = Document(path)
    assert edited.paragraphs[0].runs[0].bold is True
    assert edited.paragraphs[0].runs[1].bold is False
    inspected = server.structured_file_inspect("runs.docx")
    assert inspected["sections"][0]["header"][0] == "header text"
    assert inspected["sections"][0]["page_width_inches"] == pytest.approx(8.0, rel=0.01)
    assert inspected["sections"][0]["page_height_inches"] == pytest.approx(10.0, rel=0.01)


def test_docx_practical_editing_preserves_untouched_media_and_relationships(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("hello ").bold = True
    paragraph.add_run("world").italic = True
    linked = document.add_paragraph("link: ")
    add_hyperlink(linked, "OpenAI", "https://openai.com/")
    picture = io.BytesIO()
    Image.new("RGB", (4, 4), "red").save(picture, format="PNG")
    document.add_paragraph().add_run().add_picture(io.BytesIO(picture.getvalue()))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "a"
    table.cell(0, 1).text = "b"
    document.sections[0].header.paragraphs[0].text = "old header"
    output = io.BytesIO()
    document.save(output)
    path = root / "preserve.docx"
    path.write_bytes(output.getvalue())
    with zipfile.ZipFile(path) as package:
        media_before = {
            name: package.read(name)
            for name in package.namelist()
            if name.startswith("word/media/")
        }

    server.structured_file_apply(
        "preserve.docx",
        [
            {"op": "replace_text", "search": "hello world", "replace": "updated"},
            {"op": "table_row_insert", "table": 0, "row": 1, "values": ["x", "y"]},
            {"op": "table_column_add", "table": 0, "values": ["c", "d", "e"]},
            {"op": "table_cell_format", "table": 0, "row": 0, "column": 0, "format": {"shading": "FFFF00"}},
            {"op": "header_footer_set", "section": 0, "area": "header", "paragraph": 0, "text": "new header"},
            {"op": "style_update", "style": "Normal", "run": {"font_size_pt": 11}, "paragraph": {"space_after_pt": 3}},
        ],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )

    updated = Document(path)
    assert updated.paragraphs[0].text == "updated"
    assert len(updated.tables[0].rows) == 3
    assert len(updated.tables[0].columns) == 3
    assert updated.sections[0].header.paragraphs[0].text == "new header"
    assert any(
        relation.target_ref == "https://openai.com/"
        for relation in updated.part.rels.values()
    )
    with zipfile.ZipFile(path) as package:
        assert {
            name: package.read(name)
            for name in package.namelist()
            if name.startswith("word/media/")
        } == media_before


def test_docx_table_text_replacement_rejects_inline_relationship_loss(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    document = Document()
    table = document.add_table(rows=1, cols=1)
    add_hyperlink(table.cell(0, 0).paragraphs[0], "linked", "https://example.test/")
    output = io.BytesIO()
    document.save(output)
    path = root / "linked-cell.docx"
    path.write_bytes(output.getvalue())

    with pytest.raises(Exception, match="full-text replacement is unsupported"):
        server.structured_file_apply(
            "linked-cell.docx",
            [
                {
                    "op": "table_cell_set",
                    "table": 0,
                    "row": 0,
                    "column": 0,
                    "text": "replacement",
                }
            ],
            expected_sha256=sha256_bytes(path.read_bytes()),
        )


def test_xlsx_copy_fill_and_common_features_preserve_untouched_sheet_state(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    book = Workbook()
    sheet = book.active
    sheet.title = "Data"
    sheet.sheet_properties.tabColor = "1072BA"
    sheet.append(["Value", "Double"])
    for value in (2, 3, 4):
        sheet.append([value, None])
    sheet["B2"] = "=A2*2"
    output = io.BytesIO()
    book.save(output)
    path = root / "daily.xlsx"
    path.write_bytes(output.getvalue())

    server.structured_file_apply(
        "daily.xlsx",
        [
            {"op": "range_fill", "sheet": "Data", "source": "B2", "target": "B2:B4"},
            {"op": "freeze_panes_set", "sheet": "Data", "cell": "A2"},
            {"op": "autofilter_set", "sheet": "Data", "range": "A1:B4"},
            {"op": "table_add", "sheet": "Data", "name": "DailyData", "range": "A1:B4"},
            {"op": "validation_add", "sheet": "Data", "type": "whole", "formula1": "0", "range": "A2:A4"},
            {"op": "conditional_cell_is", "sheet": "Data", "range": "B2:B4", "operator": "greaterThan", "formula": "5", "fill": "00FF00"},
            {"op": "chart_add", "sheet": "Data", "type": "line", "data_range": "B1:B4", "anchor": "D2"},
            {"op": "page_setup_set", "sheet": "Data", "values": {"orientation": "landscape", "fitToWidth": 1}},
        ],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )
    updated = load_workbook(path, data_only=False)
    sheet = updated["Data"]
    assert [sheet[f"B{row}"].value for row in range(2, 5)] == ["=A2*2", "=A3*2", "=A4*2"]
    assert sheet.sheet_properties.tabColor.rgb.endswith("1072BA")
    assert sheet.freeze_panes == "A2"
    assert sheet.auto_filter.ref == "A1:B4"
    assert "DailyData" in sheet.tables
    assert len(sheet.data_validations.dataValidation) == 1
    assert len(sheet._charts) == 1
    assert sheet.page_setup.orientation == "landscape"
    updated.close()


def test_docx_package_patch_edits_text_without_dropping_unsupported_parts(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("hello ").bold = True
    paragraph.add_run("world").italic = True
    output = io.BytesIO()
    document.save(output)
    unsupported_part = b"bounded-placeholder"
    original = add_package_part(
        output.getvalue(), "word/vbaProject.bin", unsupported_part
    )
    path = root / "macro-carrier.docx"
    path.write_bytes(original)

    inspected = server.structured_file_inspect("macro-carrier.docx")
    assert inspected["write_rejected_features"] == ["VBA/macros"]
    assert inspected["package_patch_supported_operations"] == [
        "replace_text",
        "metadata_set",
    ]
    result = server.structured_file_apply(
        "macro-carrier.docx",
        [{"op": "replace_text", "search": "hello world", "replace": "updated"}],
        expected_sha256=sha256_bytes(original),
    )

    assert result["preservation_mode"] == "package_patch"
    assert Document(path).paragraphs[0].text == "updated"
    with zipfile.ZipFile(path) as package:
        assert package.read("word/vbaProject.bin") == unsupported_part


def test_xlsx_package_patch_updates_values_without_dropping_unsupported_parts(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    marker = b'<ext uri="urn:wlmcp-test"><feature xmlns="urn:wlmcp-test"/></ext>'
    original = replace_package_part(
        xlsx_bytes(),
        "xl/worksheets/sheet1.xml",
        lambda payload: payload.replace(
            b"</worksheet>", b"<extLst>" + marker + b"</extLst></worksheet>"
        ),
    )
    path = root / "pivot-carrier.xlsx"
    path.write_bytes(original)

    inspected = server.structured_file_inspect("pivot-carrier.xlsx")
    assert "unsupported extension lists" in inspected["write_rejected_features"]
    assert inspected["package_patch_supported_operations"] == [
        "cell_set",
        "range_set",
        "range_clear",
    ]
    result = server.structured_file_apply(
        "pivot-carrier.xlsx",
        [
            {"op": "cell_set", "sheet": "Data", "cell": "B2", "value": 7},
            {
                "op": "range_set",
                "sheet": "Data",
                "range": "C1:C2",
                "values": [["Double"], ["=B2*2"]],
            },
        ],
        expected_sha256=sha256_bytes(original),
    )

    assert result["preservation_mode"] == "package_patch"
    with pytest.warns(UserWarning, match="Unknown extension is not supported"):
        book = load_workbook(path, data_only=False)
    assert book["Data"]["B2"].value == 7
    assert book["Data"]["C1"].value == "Double"
    assert book["Data"]["C2"].value == "=B2*2"
    book.close()
    with zipfile.ZipFile(path) as package:
        assert b"urn:wlmcp-test" in package.read("xl/worksheets/sheet1.xml")

    signed = add_package_part(
        xlsx_bytes(), "_xmlsignatures/sig1.xml", b"<Signature/>"
    )
    signed_path = root / "signed.xlsx"
    signed_path.write_bytes(signed)
    with pytest.raises(ValueError, match="invalidate digital signatures"):
        server.structured_file_apply(
            "signed.xlsx",
            [{"op": "cell_set", "sheet": "Data", "cell": "B2", "value": 9}],
            expected_sha256=sha256_bytes(signed),
        )


def test_xlsx_single_cell_ranges_and_overlapping_copy_are_practical(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    book = Workbook()
    sheet = book.active
    sheet.title = "Data"
    for row, value in enumerate((1, 2, 3), start=1):
        sheet.cell(row, 1).value = value
    output = io.BytesIO()
    book.save(output)
    path = root / "ranges.xlsx"
    path.write_bytes(output.getvalue())

    server.structured_file_apply(
        "ranges.xlsx",
        [
            {"op": "range_copy", "sheet": "Data", "source": "A1:A2", "target": "A2:A3"},
            {"op": "range_clear", "sheet": "Data", "range": "A1"},
            {
                "op": "format_range",
                "sheet": "Data",
                "range": "A2",
                "format": {"font": {"bold": True}},
            },
        ],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )

    updated = load_workbook(path)
    sheet = updated["Data"]
    assert [sheet[f"A{row}"].value for row in range(1, 4)] == [None, 1, 2]
    assert sheet["A2"].font.bold is True
    updated.close()


def test_zip_preserves_directories_and_rejects_windows_collisions(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w") as output:
        output.writestr("notes/", b"")
        output.writestr("notes/a.txt", b"a")
    path = root / "paths.zip"
    path.write_bytes(archive.getvalue())

    server.structured_file_apply(
        "paths.zip",
        [{"op": "entry_add", "name": "notes/b.txt", "text": "b"}],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )
    with zipfile.ZipFile(path) as updated:
        assert "notes/" in updated.namelist()
        assert updated.read("notes/b.txt") == b"b"

    with pytest.raises(ValueError, match="already exists"):
        server.structured_file_apply(
            "paths.zip",
            [{"op": "entry_add", "name": "NOTES/A.TXT", "text": "collision"}],
            expected_sha256=sha256_bytes(path.read_bytes()),
        )
    with pytest.raises(ValueError, match="file/directory path collision"):
        server.structured_file_apply(
            "paths.zip",
            [{"op": "entry_add", "name": "notes", "text": "file"}],
            expected_sha256=sha256_bytes(path.read_bytes()),
        )


def test_image_transform_preserves_metadata_and_rejects_unsafe_cases(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    server, root = load_server(tmp_path, monkeypatch)
    image = Image.new("RGB", (32, 32), color="blue")
    exif = Image.Exif()
    exif[0x010E] = "keep me"
    path = root / "metadata.jpg"
    image.save(path, exif=exif)

    server.structured_file_apply(
        "metadata.jpg",
        [{"op": "resize", "width": 16, "height": 16}],
        expected_sha256=sha256_bytes(path.read_bytes()),
    )
    with Image.open(path) as preserved:
        assert preserved.getexif().get(0x010E) == "keep me"

    with pytest.raises(ValueError, match="image pixel count exceeds"):
        server.structured_file_apply(
            "metadata.jpg",
            [{"op": "resize", "width": 100000, "height": 100000}],
            expected_sha256=sha256_bytes(path.read_bytes()),
        )

    first = Image.new("RGB", (10, 10), color="red")
    second = Image.new("RGB", (10, 10), color="green")
    gif_path = root / "animated.gif"
    first.save(gif_path, save_all=True, append_images=[second], loop=0, duration=100)
    with pytest.raises(ValueError, match="multi-frame image transformation is unsupported"):
        server.structured_file_apply(
            "animated.gif",
            [{"op": "resize", "width": 5, "height": 5}],
            expected_sha256=sha256_bytes(gif_path.read_bytes()),
        )
